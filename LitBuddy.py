import tkinter as tk
from tkinter import ttk, messagebox
from gtts import gTTS
import pyttsx3
import os
import tempfile
import socket
import pygame
from mutagen.mp3 import MP3
import re
import time
import speech_recognition as sr
import sys
import threading
import json

last_spoken_text = ""
is_paused = False
is_highlight_paused = False
highlight_timer_id = None
loading_animating = False
loading_label = None
dyslexic_mode = False
tts_thread_started = False
after_ids = {
    "syllable_snap": None,
    "spell_game": None,
    "sound_it_out": None,
    "audio_game": None
}






# === Welcome message ===
def play_welcome():
    profiles_data = load_profiles()
    name = profiles_data.get("current_user", "Player")
    welcome_msg = f"Hi {name}!!! I'm your helper, LitBuddy! Let's learn together!!!"
    if has_internet():
        try:
            tts = gTTS(text=welcome_msg, lang="en")
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
                tts.save(tmp.name)
            os.system(f'mpg123 -q "{tmp.name}" >nul 2>&1')
            os.remove(tmp.name)
        except Exception as e:
            print("gTTS welcome failed:", e)
            if offline_engine:
                offline_engine.say(welcome_msg)
                offline_engine.runAndWait()
    elif offline_engine:
        offline_engine.say(welcome_msg)
        offline_engine.runAndWait()
    else:
        print("❌ No voice engine available.")



# === Start of Teacher Tab ===




editor_container = None


def build_phonics_editor_gui(master):
    global editor_container

    import tkinter as tk
    from tkinter import messagebox
    import tempfile
    from gtts import gTTS
    import os
    import json
    import sounddevice as sd
    from scipy.io.wavfile import write
    import numpy as np
    import pygame
    import time

    # Destroy old UI if refreshing
    if editor_container:
        editor_container.destroy()

    editor_container = tk.Frame(master, bg="#FDF6E3")
    editor_container.pack(fill="both", expand=True)

   # Load phonics map
    try:
        with open(PHONICS_MAP_PATH) as f:
            raw = json.load(f)
            if isinstance(raw, list):
                # Convert list of dicts to one dict
                data = {k: v for d in raw for k, v in d.items()}
            elif isinstance(raw, dict):
                data = raw
            else:
                print("⚠️ Unexpected format in phonics map.")
                data = {}
    except Exception as e:
        print(f"Error loading phonics map: {e}")
        data = {}


    phonics_audio_map = data

    canvas = tk.Canvas(editor_container, borderwidth=0, bg=editor_container["bg"], highlightthickness=0)
    scroll_frame = tk.Frame(canvas, bg=editor_container["bg"])
    vsb = tk.Scrollbar(editor_container, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=vsb.set)

    vsb.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)
    canvas.create_window((0, 0), window=scroll_frame, anchor="nw")

    def on_frame_configure(event):
        canvas.configure(scrollregion=canvas.bbox("all"))

    scroll_frame.bind("<Configure>", on_frame_configure)

    # Editor Grid
    tk.Label(scroll_frame, text="🔤 Grapheme", font=(CHOSEN_FONT, 10, "bold"), bg=scroll_frame["bg"]).grid(row=0, column=0, padx=5, pady=2)
    tk.Label(scroll_frame, text="🗣️ Sound or File Path", font=(CHOSEN_FONT, 10, "bold"), bg=scroll_frame["bg"]).grid(row=0, column=1, padx=5, pady=2)

    rows = []
    for i, (key, val) in enumerate(data.items(), start=1):
        g_var = tk.StringVar(value=key)
        s_var = tk.StringVar(value=val)
        tk.Entry(scroll_frame, textvariable=g_var, width=12).grid(row=i, column=0, padx=5, pady=2)
        tk.Entry(scroll_frame, textvariable=s_var, width=35).grid(row=i, column=1, padx=5, pady=2)
        rows.append((g_var, s_var))

    # Extra empty row
    g_var = tk.StringVar()
    s_var = tk.StringVar()
    row = len(rows) + 1
    tk.Entry(scroll_frame, textvariable=g_var, width=12).grid(row=row, column=0, padx=5, pady=2)
    tk.Entry(scroll_frame, textvariable=s_var, width=35).grid(row=row, column=1, padx=5, pady=2)
    rows.append((g_var, s_var))


    def load_phonics_map():
        file_path = filedialog.askopenfilename(
            title="📂 Select Phonics Map",
            filetypes=[("JSON files", "*.json")]
        )
        if not file_path:
            return

        try:
            with open(file_path) as f:
                new_data = json.load(f)
            # Replace current data
            global phonics_audio_map
            phonics_audio_map = new_data
            with open(PHONICS_MAP_PATH, "w") as f:
                json.dump(new_data, f, indent=2)
            messagebox.showinfo("✅ Loaded", f"Phonics map loaded from:\n{file_path}")
            refresh_editor()
        except Exception as e:
            messagebox.showerror("❌ Error", f"Failed to load map:\n{e}")

    def save_map():
        new_map = {}
        for g, s in rows:
            key = g.get().strip().lower()
            val = s.get().strip()
            if key:
                new_map[key] = val
        os.makedirs(os.path.dirname(PHONICS_MAP_PATH), exist_ok=True)
        with open(PHONICS_MAP_PATH, "w") as f:
            json.dump(new_map, f, indent=2)
        messagebox.showinfo("✅ Saved", "Phonics sound map updated.")

    tk.Button(scroll_frame, text="💾 Save Changes", font=(CHOSEN_FONT, 12), bg="#AED581", command=save_map).grid(row=row+1, column=0, columnspan=2, pady=10)
    tk.Button(scroll_frame, text="📂 Load Phonics Map", font=(CHOSEN_FONT, 10), command=load_phonics_map, bg="#FFE0B2").grid(row=row+1, column=2, padx=5)


    # ▶️ Chunk Tester
    chunk_entry = tk.Entry(scroll_frame, font=(CHOSEN_FONT, 12), width=10)
    chunk_entry.grid(row=row+2, column=0, padx=5, pady=5)

    def speak_chunk():
        chunk = chunk_entry.get().strip().lower()
        audio_path = phonics_audio_map.get(chunk)
        if not audio_path:
            messagebox.showwarning("Not Found", f"No entry for: {chunk}")
            return

        if os.path.isfile(audio_path):
            pygame.mixer.init()
            pygame.mixer.music.load(audio_path)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                time.sleep(0.1)
        else:
            tts = gTTS(text=audio_path)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as f:
                path = f.name
                tts.save(path)
            pygame.mixer.init()
            pygame.mixer.music.load(path)
            pygame.mixer.music.play()
            while pygame.mixer.music.get_busy():
                time.sleep(0.1)
            os.remove(path)

    tk.Button(scroll_frame, text="▶ Play", font=(CHOSEN_FONT, 10), command=speak_chunk, bg="#FFF8DC").grid(row=row+2, column=1, pady=5)

    # 🎙️ Recorder Studio
    recorder_frame = tk.LabelFrame(scroll_frame, text="🎙️ Phonogram Recorder Studio", font=(CHOSEN_FONT, 10), bg=scroll_frame["bg"])
    recorder_frame.grid(row=row+3, column=0, columnspan=2, pady=20, padx=10, sticky="ew")

    record_var = tk.StringVar()
    tk.Label(recorder_frame, text="Phonogram:", font=(CHOSEN_FONT, 10), bg=scroll_frame["bg"]).grid(row=0, column=0, padx=5)
    tk.Entry(recorder_frame, textvariable=record_var, font=(CHOSEN_FONT, 10), width=10).grid(row=0, column=1, padx=5)

    def record_audio(duration=4):
        phonogram = record_var.get().strip().lower()
        if not phonogram:
            messagebox.showwarning("⚠️ Missing", "Enter a phonogram name.")
            return

        fs = 44100
        messagebox.showinfo("⏺️ Recording", f"Recording {duration} seconds for: '{phonogram}'")
        recording = sd.rec(int(duration * fs), samplerate=fs, channels=1)
        sd.wait()

        path = f"audio/phonograms/{phonogram}.wav"
        os.makedirs(os.path.dirname(path), exist_ok=True)
        write(path, fs, recording)

        messagebox.showinfo("✅ Saved", f"Recording saved to {path}")

    def play_recording():
        phonogram = record_var.get().strip().lower()
        path = f"audio/phonograms/{phonogram}.wav"
        if os.path.isfile(path):
            pygame.mixer.init()
            pygame.mixer.music.load(path)
            pygame.mixer.music.play()
        else:
            messagebox.showwarning("⚠️ Not Found", f"No recording found for '{phonogram}'.")

    def delete_recording():
        phonogram = record_var.get().strip().lower()
        path = f"audio/phonograms/{phonogram}.wav"
        if os.path.isfile(path):
            os.remove(path)
            messagebox.showinfo("🗑️ Deleted", f"Deleted: {path}")
        else:
            messagebox.showwarning("⚠️ Missing", f"No file to delete.")

    def add_to_phonics_map():
        phonogram = record_var.get().strip().lower()
        path = f"audio/phonograms/{phonogram}.wav"
        if os.path.isfile(path):
            data[phonogram] = path
            with open(PHONICS_MAP_PATH, "w") as f:
                json.dump(data, f, indent=2)
            messagebox.showinfo("✅ Added", f"'{phonogram}' mapped to: {path}")
        else:
            messagebox.showwarning("⚠️ Not Found", f"Record file first.")
    


    def refresh_editor():
        build_phonics_editor_gui(master)

    



   


    # Buttons Row

    tk.Button(recorder_frame, text="🎤 Record", font=(CHOSEN_FONT, 10), command=record_audio, bg="#E1F5FE").grid(row=0, column=2, padx=4)
    tk.Button(recorder_frame, text="▶ Play", font=(CHOSEN_FONT, 10), command=play_recording, bg="#DCEDC8").grid(row=0, column=3, padx=4)
    tk.Button(recorder_frame, text="🗑 Delete", font=(CHOSEN_FONT, 10), command=delete_recording, bg="#FFCDD2").grid(row=0, column=4, padx=4)
    tk.Button(recorder_frame, text="➕ Add to Map", font=(CHOSEN_FONT, 10), command=add_to_phonics_map, bg="#D1C4E9").grid(row=0, column=5, padx=4)
    tk.Button(recorder_frame, text="🔄 Refresh", font=(CHOSEN_FONT, 10), command=refresh_editor, bg="#FFF59D").grid(row=0, column=6, padx=4)

    from audio_engine import generate_syllable_audio

    from tkinter import filedialog

    def run_syllable_audio_generator():
        file_path = filedialog.askopenfilename(
            title="📂 Select Wordlist File",
            filetypes=[("Text files", "*.txt"), ("JSON files", "*.json")]
        )
        if not file_path:
            return
        try:
            generate_syllable_audio(file_path)
            messagebox.showinfo("✅ Done", f"Syllable audio generated for:\n{os.path.basename(file_path)}")
        except Exception as e:
            messagebox.showerror("❌ Error", f"Failed to generate audio:\n{e}")

    tk.Button(scroll_frame, text="🎧 Generate Syllable Audio from Wordlist",
              font=(CHOSEN_FONT, 12), bg="#FFF9C4",
              command=run_syllable_audio_generator).grid(row=row+4, column=0, columnspan=2, pady=10)










def record_audio(duration=4):
    phonogram = record_var.get().strip().lower()
    if not phonogram:
        tk.messagebox.showwarning("⚠️ Missing", "Enter a phonogram to name the file.")
        return

    fs = 44100  # Sample rate
    tk.messagebox.showinfo("⏺️ Recording", f"Recording 4 seconds for: {phonogram}")
    recording = sd.rec(int(duration * fs), samplerate=fs, channels=1)
    sd.wait()

    path = f"audio/phonograms/{phonogram}.wav"
    os.makedirs(os.path.dirname(path), exist_ok=True)
    write(path, fs, recording)

    # Update phonics_audio_map.json
    data[phonogram] = path
    with open(PHONICS_MAP_PATH, "w") as f:
        json.dump(data, f, indent=2)

    tk.messagebox.showinfo("✅ Saved", f"Recording for '{phonogram}' saved.\nUpdated phonics map.")

    tk.Button(recorder_frame, text="🎤 Record", font=(CHOSEN_FONT, 10), command=record_audio, bg="#E1F5FE").grid(row=0, column=2, padx=10)





def build_phonics_converter_gui(parent_frame):
    import tkinter as tk
    from tkinter import filedialog, messagebox
    import os, csv, json
    import openpyxl
    import docx
    import pronouncing

    phoneme_var = tk.StringVar()


    status_label = tk.Label(parent_frame, text="", font=(CHOSEN_FONT, 10), bg=parent_frame["bg"])
    status_label.pack(pady=10)

    phonics_map = {
        "h": {"sound": "h", "example": "hat"},
        "ʌ": {"sound": "short u", "example": "cup"},
        "z": {"sound": "z", "example": "zoo"},
        "b": {"sound": "b", "example": "bat"},
        "n": {"sound": "n", "example": "net"},
        "d": {"sound": "d", "example": "dog"},
        "ɛ": {"sound": "short e", "example": "bed"},
        "l": {"sound": "l", "example": "lamp"},
        "f": {"sound": "f", "example": "fish"},
        "t": {"sound": "t", "example": "top"},
        "j": {"sound": "y", "example": "yellow"},
        "oʊ": {"sound": "long o", "example": "go"},
        "tʃ": {"sound": "ch", "example": "chicken"},
        "ɪ": {"sound": "short i", "example": "sit"},
        "k": {"sound": "k", "example": "cat"},
        "æ": {"sound": "short a", "example": "cat"},
        "s": {"sound": "s", "example": "sun"},
        "u": {"sound": "long u", "example": "flute"}
    }


    def get_phonics_ready_words(wordlist, phonics_map):
        ready_words = {}
        for word, data in wordlist.items():
            phonemes = data["ipa"].split()  # This gives clean phonemes
            if all(p in phonics_map for p in phonemes):
                ready_words[word] = {
                    **data,
                    "phonemes": phonemes,
                    "phonics": [phonics_map[p] for p in phonemes]
                }
        return ready_words



    def save_phonics_ready_words(ready_words, filename="phonics_ready.json"):
        filepath = WORDLIST_DIR / filename
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(ready_words, f, indent=2, ensure_ascii=False)

    def filter_phonics_ready_words():
        try:
            # Load the IPA wordlist
            with open("syllables_with_ipa.json", "r", encoding="utf-8") as f:
                wordlist = json.load(f)

            # Filter and save
            ready_words = get_phonics_ready_words(wordlist, phonics_map)
            save_phonics_ready_words(ready_words)

            status_label.config(
                text=f"✅ {len(ready_words)} phonics-ready words saved to phonics_ready.json",
                fg="#2E7D32"
            )
        except Exception as e:
            status_label.config(text=f"❌ Error: {e}", fg="#D32F2F")

    tk.Button(parent_frame, text="🧠 Filter Phonics-Ready Words from IPA JSON",
          font=(CHOSEN_FONT, 12), bg="#C8E6C9", command=filter_phonics_ready_words).pack(pady=5)






    def simple_grapheme_split(word):
        graphemes = ["ch", "sh", "th", "wh", "ph", "ck", "qu", "ng", "ee", "oo", "ay", "ou", "ow", "ar", "or"]
        result = []
        i = 0
        while i < len(word):
            chunk = word[i:i+2]
            if chunk in graphemes:
                result.append(chunk)
                i += 2
            else:
                result.append(word[i])
                i += 1
        return result

    def read_words(file_path):
        ext = os.path.splitext(file_path)[1].lower()
        words = []

        if ext == ".txt":
            with open(file_path) as f:
                words = [line.strip() for line in f if line.strip()]
        elif ext == ".csv":
            with open(file_path, newline='') as f:
                reader = csv.DictReader(f)
                if "word" in reader.fieldnames:
                    words = [row["word"].strip() for row in reader if row.get("word")]
                else:
                    f.seek(0)
                    reader = csv.reader(f)
                    words = [row[0].strip() for row in reader if row and row[0]]
        elif ext == ".xlsx":
            wb = openpyxl.load_workbook(file_path)
            sheet = wb.active
            col = 1
            words = [str(cell.value).strip() for cell in sheet.iter_rows(min_row=2, min_col=col, max_col=col, values_only=True) if cell[0]]
        elif ext == ".docx":
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                for line in para.text.splitlines():
                    line = line.strip()
                    if line:
                        words.append(line)
        elif ext == ".json":
            with open(file_path) as f:
                data = json.load(f)
                words = [w["word"].strip() for w in data if "word" in w]
        else:
            raise ValueError("Unsupported file format.")
        
        return list(set(w.lower() for w in words if w))

    def convert(file_path):
        try:
            word_list = read_words(file_path)
        except Exception as e:
            status_label.config(text=f"❌ Error: {e}", fg="#D32F2F")
            return

        result = []
        for word in word_list:
            phonics = simple_grapheme_split(word)
            phonemes = pronouncing.phones_for_word(word)
            entry = {
                "word": word,
                "phonics": phonics,
                "phonemes": phonemes[0].split() if phonemes else []
            }
            result.append(entry)

        out_path = os.path.splitext(file_path)[0] + "_with_phonics.json"
        with open(out_path, "w") as out:
            json.dump(result, out, indent=2)

        status_label.config(
            text=f"✅ {len(result)} words converted.\nSaved as {os.path.basename(out_path)}",
            fg="#2E7D32"
        )

    def select_and_convert():
        file_path = filedialog.askopenfilename(
            title="Select Wordlist File",
            filetypes=[
                ("All Supported", "*.txt *.csv *.xlsx *.docx *.json"),
                ("Text Files", "*.txt"),
                ("CSV Files", "*.csv"),
                ("Excel Files", "*.xlsx"),
                ("Word Docs", "*.docx"),
                ("JSON", "*.json")
            ]
        )
        if file_path:
            convert(file_path)

    convert_btn = tk.Button(parent_frame, text="📂 Select & Convert Wordlist (Phonics Only)",
                            font=(CHOSEN_FONT, 12), bg="#FFF8E1", command=select_and_convert)
    convert_btn.pack(pady=10)

    # Optional Auto Converter (placeholder)
    def auto_convert_to_phonics_file():
        messagebox.showinfo("Coming Soon", "Auto-convert isn't configured in this simplified version.")

    tk.Button(parent_frame, text="🔁 Auto-Convert to Phonics JSON",
              font=(CHOSEN_FONT, 12), bg="#D1C4E9", command=auto_convert_to_phonics_file).pack(pady=5)






def auto_guess_phonics(word):
    word = word.lower()
    phonograms = [
        # Multiletter phonograms
        "ough", "augh", "eigh", "igh", "dge", "tch", "tion", "sion", "ph", "gn", "kn", "wr", "wh", "sh", "ch", "th",
        # Vowel teams
        "ai", "ay", "ee", "ea", "ie", "ei", "ey", "oe", "oa", "ou", "ow", "ue", "ew",
        # R-controlled vowels
        "ar", "er", "ir", "or", "ur",
        # Common blends
        "bl", "cl", "fl", "gl", "pl", "sl", "br", "cr", "dr", "fr", "gr", "pr", "tr", "sc", "sk", "sm", "sn", "sp", "st", "sw", "tw", "spr", "str", "spl", "scr", "shr", "thr",
        # Common suffixes & endings
        "ed", "ing", "er", "est", "ful", "less", "ness", "ment", "ly", "y",
        # Short vowel fallback (handled implicitly)
    ]

    chunks = []
    i = 0
    while i < len(word):
        matched = False
        for p in sorted(phonograms, key=len, reverse=True):
            if word[i:i+len(p)] == p:
                chunks.append(p)
                i += len(p)
                matched = True
                break
        if not matched:
            chunks.append(word[i])
            i += 1
    return chunks





    convert_btn = tk.Button(parent_frame, text="📂 Select & Convert Wordlist (Phonics Only)", font=(CHOSEN_FONT, 12),bg="#FFF8E1", command=select_and_convert)
    convert_btn.pack(pady=10)
    tk.Button(parent_frame, text="🔁 Auto-Convert to Phonics JSON", font=(CHOSEN_FONT, 12), bg="#D1C4E9", command=auto_convert_to_phonics_file).pack(pady=5)
    






def build_word_fixer_gui(parent_frame):
    import tkinter as tk
    from tkinter import filedialog, messagebox
    import json
    import os
    from phonemizer import phonemize
    from playsound import playsound
    import subprocess
    import pygame


    file_data = []
    word_list = []
    filtered_list = []
    selected_word = tk.StringVar()
    file_path = tk.StringVar(value="No file selected")
    search_var = tk.StringVar()
    phoneme_var = tk.StringVar()



    def refresh_dropdown():
        dropdown["menu"].delete(0, "end")
        for w in filtered_list:
            dropdown["menu"].add_command(label=w, command=tk._setit(selected_word, w, display_selected_word))
        if filtered_list:
            selected_word.set(filtered_list[0])
            display_selected_word()
        else:
            phonics_var.set("")
            blend_var.set("")
            type_var.set("auto")
            selected_word.set("")

    # Header
    tk.Label(parent_frame, text="🛠️ Fix Phonics for Word", font=(CHOSEN_FONT, 14, "bold"),
             bg=parent_frame["bg"]).pack(pady=(10, 5))



    def play_audio(audio_path):
        pygame.mixer.init()
        pygame.mixer.music.load(audio_path)
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)


    def get_phonemes(word):
        return phonemize(
            word,
            language="en-us",
            backend="espeak",
            strip=True,
            preserve_punctuation=True,
            with_stress=False
        ).split()

    '''
    def synthesize_phoneme(phoneme, output_dir="phonemes_audio"):
        os.makedirs(output_dir, exist_ok=True)
        safe_name = phoneme.replace("ʃ", "sh").replace("ɪ", "ih").replace("ʧ", "ch")  # customize as needed
        filename = os.path.join(output_dir, f"{safe_name}.wav")
        if not os.path.exists(filename):
            subprocess.run(["espeak", "-v", "en-us", "--ipa", f"[{phoneme}]", "-w", filename])
        return filename
    '''

    def synthesize_phoneme_audio(phoneme, output_dir="phonemes_audio"):
        os.makedirs(output_dir, exist_ok=True)
        filename = os.path.join(output_dir, f"{phoneme}.wav")
        if not os.path.exists(filename):
            try:
                subprocess.run(["espeak-ng", "-w", filename, phoneme])
            except Exception as e:
                print(f"Error synthesizing {phoneme}: {e}")
        return filename


    



    def synthesize_and_play_ipa(ipa_list, output_dir="phonemes_audio"):
        os.makedirs(output_dir, exist_ok=True)

        for phoneme in ipa_list:
            safe_name = phoneme.replace("ʃ", "sh").replace("ɪ", "ih").replace("ʧ", "ch").replace("ʤ", "j").replace("θ", "th")
            filename = os.path.join(output_dir, f"{safe_name}.wav")

            if not os.path.exists(filename):
                try:
                    subprocess.run(["espeak", "-v", "en-us", "--ipa", f"[{phoneme}]", "-w", filename])
                except Exception as e:
                    print(f"Error synthesizing {phoneme}: {e}")
                    continue

            try:
                playsound(filename)
            except Exception as e:
                print(f"Error playing {filename}: {e}")


    def play_phonemes(word):
        phonemes = get_phonemes(word)
        for p in phonemes:
            audio_path = synthesize_phoneme(p)
            play_audio(audio_path)

    def play_phonics_chunks():
        chunks = [c.strip() for c in phonics_var.get().split(",") if c.strip()]
        ipa_list = phonics_chunks_to_ipa(chunks)
        synthesize_and_play_ipa(ipa_list)


    grapheme_to_ipa = {
        "sh": "ʃ",
        "ch": "ʧ",
        "th": "θ",
        "ng": "ŋ",
        "ph": "f",
        "wh": "w",
        "qu": "kw",
        "ck": "k",
        "ee": "iː",
        "oo": "uː",
        "ar": "ɑː",
        "or": "ɔː",
        "igh": "aɪ",
        "ow": "aʊ",
        "ay": "eɪ",
        "ai": "eɪ",
        "ea": "iː",
        "ie": "aɪ",
        "ou": "aʊ",
        "oy": "ɔɪ",
        "oi": "ɔɪ",
        "ur": "ɜː",
        "er": "ɜː",
        "ir": "ɜː",
        "eigh": "eɪ",
        "dge": "ʤ",
        "tch": "ʧ",
        "y": "j",
        # Add more as needed
    }


    def phonics_chunks_to_ipa(chunks):
        ipa_list = []
        for chunk in chunks:
            ipa = grapheme_to_ipa.get(chunk.lower(), chunk.lower())
            ipa_list.append(ipa)
        return ipa_list



    


    # File selection
    def load_wordlist():
        path = filedialog.askopenfilename(
            title="Open wordlist JSON",
            filetypes=[("JSON files", "*.json")]
        )
        if path:
            file_path.set(path)
            try:
                with open(path) as f:
                    nonlocal file_data, word_list, filtered_list
                    file_data = json.load(f)
                    word_list = [entry.get("word", "") for entry in file_data]
                    filtered_list = word_list.copy()
                    search_var.set("")
                    refresh_dropdown()
                    status_label.config(text=f"✅ Loaded {len(word_list)} words", fg="#2E7D32")
            except Exception as e:
                messagebox.showerror("Error", f"Couldn't load file:\n{e}")

    # File UI
    file_frame = tk.Frame(parent_frame, bg=parent_frame["bg"])
    file_frame.pack(pady=5)
    tk.Button(file_frame, text="📂 Select Wordlist", font=(CHOSEN_FONT, 10), command=load_wordlist).pack(side="left", padx=5)
    tk.Label(file_frame, textvariable=file_path, font=(CHOSEN_FONT, 8), bg=parent_frame["bg"]).pack(side="left", padx=5)

    status_label = tk.Label(parent_frame, text="", font=(CHOSEN_FONT, 10), bg=parent_frame["bg"])
    status_label.pack()

    # Search bar
    search_entry = tk.Entry(parent_frame, textvariable=search_var, font=(CHOSEN_FONT, 11), width=30)
    search_entry.pack(pady=(5, 2))
    tk.Label(parent_frame, text="🔍 Start typing to filter words", font=(CHOSEN_FONT, 9), bg=parent_frame["bg"], fg="#777").pack()

    # Dropdown
    dropdown = tk.OptionMenu(parent_frame, selected_word, "")
    dropdown.config(font=(CHOSEN_FONT, 12), width=25)
    dropdown.pack(pady=5)

    # Word fields
    phonics_var = tk.StringVar()
    blend_var = tk.StringVar()
    type_var = tk.StringVar(value="auto")

    tk.Label(parent_frame, text="Phonics Chunks (comma-separated):", font=(CHOSEN_FONT, 10), bg=parent_frame["bg"]).pack(pady=(10, 0))
    tk.Entry(parent_frame, textvariable=phonics_var, width=50).pack()

    tk.Label(parent_frame, text="Blend Phrase (optional):", font=(CHOSEN_FONT, 10), bg=parent_frame["bg"]).pack(pady=(10, 0))
    tk.Entry(parent_frame, textvariable=blend_var, width=50).pack()

    tk.Label(parent_frame, text="Phonics Mode:", font=(CHOSEN_FONT, 10), bg=parent_frame["bg"]).pack(pady=(10, 0))
    tk.OptionMenu(parent_frame, type_var, "auto", "custom").pack()
    tk.Label(parent_frame, text="IPA Phonemes:", font=(CHOSEN_FONT, 10), bg=parent_frame["bg"]).pack(pady=(10, 0))
    tk.Label(parent_frame, textvariable=phoneme_var, font=(CHOSEN_FONT, 11), bg=parent_frame["bg"], fg="#444").pack()


    '''
    def display_selected_word(*args):
        word = selected_word.get()
        for entry in file_data:
            if entry.get("word") == word:
                phonics_var.set(", ".join(entry.get("phonics", [])))
                blend_var.set(entry.get("blend_phrase", ""))
                type_var.set(entry.get("phonics_type", "auto"))
                break
    '''
    def display_selected_word(*args):
        word = selected_word.get()
        for entry in file_data:
            if entry.get("word") == word:
                phonics_var.set(", ".join(entry.get("phonics", [])))
                blend_var.set(entry.get("blend_phrase", ""))
                type_var.set(entry.get("phonics_type", "auto"))
                break
        phoneme_var.set(" ".join(get_phonemes(word)))


    def save_word_update():
        word = selected_word.get()
        for entry in file_data:
            if entry.get("word") == word:
                entry["phonics"] = [p.strip() for p in phonics_var.get().split(",") if p.strip()]
                entry["blend_phrase"] = blend_var.get().strip()
                entry["phonics_type"] = type_var.get().strip()
                break
        try:
            with open(file_path.get(), "w") as f:
                json.dump(file_data, f, indent=2)
            status_label.config(text=f"💾 Saved updates to {os.path.basename(file_path.get())}", fg="#33691E")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save file:\n{e}")

    

    button_frame = tk.Frame(parent_frame)
    button_frame.pack(pady=10)

    tk.Button(button_frame, text="💾 Save Updates", font=(CHOSEN_FONT, 12), bg="#AED581", command=save_word_update).pack(side="left", padx=5)
    tk.Button(button_frame, text="🔊 Play Phonemes", font=(CHOSEN_FONT, 12), bg="#FFCCBC", command=lambda: play_phonemes(selected_word.get())).pack(side="left", padx=5)
    tk.Button(button_frame, text="🔊 Play Phonics Chunks", font=(CHOSEN_FONT, 12), bg="#E1F5FE", command=play_phonics_chunks).pack(side="left", padx=5)








    # Handle live search
    def on_search_change(*args):
        query = search_var.get().lower()
        nonlocal filtered_list
        filtered_list = [w for w in word_list if query in w.lower()]
        refresh_dropdown()

    search_var.trace_add("write", on_search_change)








# === End of Teacher Tab ===

#===========================


# === Start of Profiles ===

PROFILE_FILE = "profiles/litbuddy.json"
os.makedirs("profiles", exist_ok=True)

PHONICS_MAP_PATH = "config/phonics_audio_map.json"



def load_profiles():
    if not os.path.exists(PROFILE_FILE):
        return {"current_user": None, "profiles": {}}
    with open(PROFILE_FILE, "r") as f:
        return json.load(f)

def refresh_user_labels():
    profiles_data = load_profiles()
    user = profiles_data.get("current_user", "Player")
    game_label.config(text=f"🎮 {user}'s Game Menu")
    name_label.config(text=f"👤 Logged in as: {user}")

def save_profiles(data):
    with open(PROFILE_FILE, "w") as f:
        json.dump(data, f, indent=2)



def reload_user_profile():
    global config, spell_words

    profiles_data = load_profiles()
    user = profiles_data.get("current_user")

    if not user:
        print("⚠️ No active user.")
        return

    user_data = profiles_data["profiles"].get(user)
    if not user_data:
        print(f"⚠️ Profile for '{user}' not found.")
        return

    # Load their word list
    selected_list = os.path.basename(user_data.get("wordlist", "default.json"))
    config["active_list"] = selected_list


    # Load words from the file
    spell_words = load_spell_words()

    print(f"✅ Loaded profile: {user} using '{selected_list}'")


def open_profile_picker():
    picker = tk.Toplevel()
    picker.title("👤 Choose Your Profile")

    tk.Label(picker, text="Select your name:", font=(CHOSEN_FONT, 12)).pack(pady=10)

    listbox = tk.Listbox(picker, height=6, width=30)
    listbox.pack(pady=5)

    # Load current profiles
    profiles_data = load_profiles()
    for name in profiles_data["profiles"].keys():
        listbox.insert(tk.END, name)

    def use_selected():
        selection = listbox.curselection()
        if selection:
            selected_name = listbox.get(selection[0])
            profiles_data["current_user"] = selected_name
            save_profiles(profiles_data)
            play_welcome()  # 👋 Personalized welcome here
            picker.destroy()
            reload_user_profile()
            refresh_user_labels()
            game_label.config(text=f"🎮 {selected_name}'s Game Menu")
            name_label.config(text=f"👤 Logged in as: {selected_name}")

    def add_new_profile():
        name = simpledialog.askstring("➕ New Profile", "Enter name:")
        if name:
            name = name.strip()
            if name in profiles_data["profiles"]:
                messagebox.showerror("⚠️ Exists", "This profile already exists.")
                return
            profiles_data["profiles"][name] = {
                "wordlist": "custom_words.json",
                "progress": {"games_played": 0, "correct": 0, "incorrect": 0}
            }
            profiles_data["current_user"] = name
            save_profiles(profiles_data)
            listbox.insert(tk.END, name)
            messagebox.showinfo("✅ Added", f"Profile '{name}' created.")
            play_welcome()

            picker.destroy()
            reload_user_profile()


    def delete_selected():
        selection = listbox.curselection()
        if not selection:
            return

        selected_name = listbox.get(selection[0])

        # 🔐 Require admin password
        if not check_password():
            messagebox.showerror("Access Denied", "Incorrect admin password.")
            return

        if messagebox.askyesno("⚠️ Confirm", f"Delete profile '{selected_name}'?"):
            profiles_data = load_profiles()

            if len(profiles_data["profiles"]) <= 1:
                messagebox.showwarning("⚠️ Required", "At least one profile must remain.")
                return

            if profiles_data["current_user"] == selected_name:
                profiles_data["current_user"] = None

            profiles_data["profiles"].pop(selected_name, None)
            save_profiles(profiles_data)
            listbox.delete(selection[0])

            messagebox.showinfo("🗑️ Deleted", f"Profile '{selected_name}' removed.")



        


    tk.Button(picker, text="➕ Add New", command=add_new_profile).pack(pady=5)
    tk.Button(picker, text="🚀 Use Selected", command=use_selected).pack(pady=(0, 10))
    tk.Button(picker, text="🗑️ Delete Selected", command=delete_selected).pack(pady=(0, 10))

   


# === End of Profiles ===





def show_loading_spinner():
    global loading_label, loading_animating
    loading_animating = True

    parent = text_box.master  # Works even if we rename the tab in the future
    loading_label = tk.Label(parent, text="📚 getting ready to read", font=(CHOSEN_FONT, 11), bg=parent["bg"])
    loading_label.pack(pady=10)
    parent.update_idletasks()

    dots = ["", ".", "..", "..."]
    def animate(i=0):
        if not loading_animating: return
        loading_label.config(text="📚 getting ready to read" + dots[i % 4])
        parent.after(100, lambda: animate(i + 1))
    animate()


def get_gtts_audio(text):
    try:
        from gtts import gTTS
        filename = "temp_audio.mp3"

        # Check selected style
        pace = read_style.get() if read_style else "standard"

        # Simulate a slower reading style by modifying punctuation pacing
        if pace == "slow":
            # Add extra pauses between sentences (gTTS respects punctuation)
            text = text.replace(". ", "... ").replace("! ", "...! ").replace("? ", "...? ")

        tts = gTTS(text=text, lang="en", slow=False)
        tts.save(filename)
        return filename

    except Exception as e:
        print("TTS error:", e)
        return None



# === Check for internet ===
def has_internet(host="8.8.8.8", port=53, timeout=2):
    try:
        socket.setdefaulttimeout(timeout)
        socket.socket(socket.AF_INET, socket.SOCK_STREAM).connect((host, port))
        return True
    except socket.error:
        return False


# === Check if OpenDyslexic font is installed ===
def preferred_font():
    try:
        test = tk.Tk()
        test_font = ("OpenDyslexic", 12)
        test_label = tk.Label(test, text="test", font=test_font)
        test_label.pack()
        test.update_idletasks()
        actual_font = test_label.cget("font")
        test.destroy()
        if "OpenDyslexic" in actual_font:
            return "OpenDyslexic"
    except:
        pass
    return "Comic Sans MS"

CHOSEN_FONT = preferred_font()


# === Fallback voice engine setup ===
def init_offline_tts():
    try:
        engine = pyttsx3.init()
        engine.setProperty('rate', 150)
        for voice in engine.getProperty('voices'):
            if "english" in voice.name.lower():
                engine.setProperty('voice', voice.id)
                break
        return engine
    except Exception as e:
        print("Offline voice init error:", e)
        return None

offline_engine = init_offline_tts()


#*******THIS IS THE START OF THE LISTEN FUNCTION**********
loading_label = None


def try_remove(path):
    try:
        os.remove(path)
    except FileNotFoundError:
        print(f"🔍 File already deleted: {path}")

def cleanup_after_playback(path):
    try_remove(path)
    enable_replay()




def cleanup_audio(path):
    try:
        os.remove(path)
    except Exception as e:
        print("Failed to remove audio:", e)




def toggle_settings():
    if settings_visible.get():
        settings_frame.place_forget()
        toggle_btn.config(text="⚙️ Show Voice Settings")
        settings_visible.set(False)
    else:
        settings_frame.place(relx=0.5, rely=0.15, anchor="n")
        toggle_btn.config(text="⚙️ Hide Voice Settings")
        settings_visible.set(True)



def preview_voice():
    sample = "This is how your voice will sound."
    engine = pyttsx3.init()
    engine.setProperty('rate', rate_slider.get())
    engine.say(sample)
    engine.runAndWait()


def split_into_sentences(text):
    return [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s]

# === Speak text entered by child ===
pygame.mixer.init()


def speak_text():
    cancel_highlight()
    global last_spoken_text
    text = text_box.get("1.0", tk.END).strip()
    if not text:
        messagebox.showinfo("Empty Box", "Let’s write something first!")
        return

    last_spoken_text = text

    # Show loading message
    loading_label = tk.Label(root, text="🎤 Preparing voice...", font=(CHOSEN_FONT, 12), fg="#555", bg="#F0F8FF")
    loading_label.pack()
    root.update_idletasks()

    if has_internet():        
        audio_path = get_gtts_audio(text)
        print("✅ Audio Path:", audio_path)
        if not audio_path:
            loading_label.destroy()
            #fallback_offline_tts(text)
            speak_with_gtts(text) #delete this

            return

        try:
            loading_label.destroy()

            # Break into sentences and get positions
            sentences = split_into_sentences(text)
            sentence_map = get_sentence_positions(sentences, text)

            if not sentence_map:
                raise Exception("Could not map sentence positions.")

            print("📣 Speak text triggered") #delete this




            # Estimate duration per sentence
            audio_duration = MP3(audio_path).info.length
            #interval = int((audio_duration / len(sentence_map)) * 1000 * 1.15)
            intervals = get_hybrid_intervals(sentences, audio_duration)
            root.after(300, lambda: highlight_sentences(sentence_map, intervals))


            # Play audio
            pygame.mixer.music.load(audio_path)
            disable_replay()
            pygame.mixer.music.play()

            # Start sentence-based highlighting
            root.after(50, lambda: wait_for_playback_and_start(sentence_map, intervals))

            # Schedule cleanup
            #root.after(int(audio_duration * 1000) + 500, lambda *_: os.remove(audio_path), enable_replay)
            root.after(int(audio_duration * 1000) + 500, lambda *_: cleanup_after_playback(audio_path))


        except Exception as e:
            print("TTS playback or sync failed:", e)
            fallback_offline_tts(text)
    else:
        loading_label.destroy()
        fallback_offline_tts(text)


def wait_for_playback_and_start(positions, intervals):
    if pygame.mixer.music.get_busy():
        highlight_sentences(positions, intervals)
    else:
        root.after(50, lambda: wait_for_playback_and_start(positions, intervals))



def fallback_offline_tts(text):
    print(f"🎙 Reading: {text}") #delete this

    if offline_engine:
        try:
            offline_engine.say(text)
            offline_engine.runAndWait()
        except Exception as e:
            print("Offline TTS error:", e)
            messagebox.showerror("Voice Error", "Oops! No voice available.")
    else:
        messagebox.showerror("Voice Error", "No speech system available.")


def replay_last_text():
    text = text_box.get("1.0", tk.END).strip()
    if not text:
        messagebox.showinfo("Replay", "There's nothing to replay yet!")
        return

    global last_spoken_text
    last_spoken_text = text

    stop_audio()         # <- clean the slate
    cancel_highlight()   # <- double-check visual reset
    speak_text_from_text(text)  # restart fresh



def speak_text_from_text(text):
    stop_audio()
    cancel_highlight()
    


    try:
        audio_path = get_gtts_audio(text)
        if not audio_path:
            fallback_offline_tts(text)
            return

        sentences = split_into_sentences(text)
        sentence_map = get_sentence_positions(sentences, text)
        audio_duration = MP3(audio_path).info.length
        intervals = get_hybrid_intervals(sentences, audio_duration)


        show_loading()
        pygame.mixer.music.load(audio_path)
        disable_replay()
        hide_loading_spinner()
        pygame.mixer.music.play()

        root.after(300, lambda: wait_for_playback_and_start(sentence_map, intervals))        
        root.after(int(audio_duration * 1000) + 500, lambda *_: [try_remove(audio_path), enable_replay()])



    except Exception as e:
        print("Replay error:", e)





def read_from_cursor():    
    stop_audio()
    cancel_highlight()
    index = text_box.index(tk.CURRENT)

    # If CURRENT is empty, fallback to 'insert' or selection start
    if not index or index == "":
        index = text_box.index("insert")

    # Prevent the bug where count("1.0", "1.0") returns nothing
    if index == "1.0":
        offset = 0
    else:
        try:
            counts = text_box.count("1.0", index, "chars")
            offset = int(counts[0]) if counts else 0
        except Exception as e:
            print(f"⚠️ Could not determine offset from index {index}: {e}")
            offset = 0

    # Visually scroll to the selected spot
    text_box.see(index)

    # Optional: flicker highlight
    def flash_start():
        text_box.tag_remove("flash_start", "1.0", tk.END)
        text_box.tag_add("flash_start", index)
        text_box.tag_configure("flash_start", background="#FFEB3B")
        root.after(300, lambda: text_box.tag_remove("flash_start", "1.0", tk.END))
    flash_start()

    # Grab text from current cursor to end
    text = text_box.get(index, tk.END).strip()
    full_text = text_box.get("1.0", tk.END)

    if not text:
        print("🔕 No text found from this point onward.")
        return

    global last_spoken_text
    last_spoken_text = text

    

    try:
        audio_path = get_gtts_audio(text)
        if not audio_path:
            fallback_offline_tts(text)
            return

        sentences = split_into_sentences(text)
        sentence_map = get_sentence_positions(sentences, full_text, offset=offset)
        audio_duration = MP3(audio_path).info.length
        intervals = get_hybrid_intervals(sentences, audio_duration)
     
        # Show spinner first
        show_loading_spinner()

        # Let the UI actually update BEFORE gTTS blocks
        root.after(100, lambda: generate_and_continue(text, full_text, offset))

        #pygame.mixer.music.play()

        root.after(300, lambda: wait_for_playback_and_start(sentence_map, intervals))
        root.after(int(audio_duration * 1000) + 500, lambda *_: cleanup_after_playback(audio_path))

        disable_replay()

    except Exception as e:
        print(f"Read from cursor error: {e}")
        fallback_offline_tts(text)




def show_context_menu(event):
    context_menu.tk_popup(event.x_root, event.y_root)



def pause_audio():
    global is_paused, is_highlight_paused
    if pygame.mixer.music.get_busy():
        pygame.mixer.music.pause()
        is_paused = True
        is_highlight_paused = True


def resume_audio():
    global is_paused, is_highlight_paused
    if is_paused:
        pygame.mixer.music.unpause()
        is_paused = False
        is_highlight_paused = False



def stop_audio():
    global is_paused, is_highlight_paused, highlight_timer_id
    pygame.mixer.music.stop()
    is_paused = False
    is_highlight_paused = False

    if highlight_timer_id is not None:
        root.after_cancel(highlight_timer_id)
        highlight_timer_id = None

    text_box.tag_remove("sentence_highlight", "1.0", tk.END)
    margin_canvas.delete("ball")
    enable_replay()




def disable_replay():
    replay_btn.config(state="disabled", text="⏳ Reading...")
    replay_btn.update_idletasks()


def enable_replay():
    replay_btn.config(state="normal",text="🔁 Replay")
    replay_btn.update_idletasks()


def split_into_sentences(text):
    return [s.strip() for s in re.split(r'(?<=[.!?]) +', text) if s]


def get_sentence_positions(sentences, full_text, offset=0):
    positions = []
    current_offset = offset

    for s in sentences:
        start_char = full_text.find(s, current_offset)
        if start_char == -1:
            print(f"⚠️ Could not match sentence: {s}")
            continue

        end_char = start_char + len(s)
        start_index = f"1.0 + {start_char}c"
        end_index = f"1.0 + {end_char}c"
        positions.append((start_index, end_index))

        current_offset = end_char  # advance search cursor forward

    return positions


def highlight_sentences(positions, intervals, index=0):
    global is_highlight_paused, highlight_timer_id

    # Stop if we're done
    if index >= len(positions):
        text_box.tag_remove("sentence_highlight", "1.0", tk.END)
        margin_canvas.delete("ball")
        highlight_timer_id = None
        return

    # If paused, try again later WITHOUT stacking timers
    if is_highlight_paused:
        if highlight_timer_id:
            root.after_cancel(highlight_timer_id)
        highlight_timer_id = root.after(300, lambda: highlight_sentences(positions, intervals, index))
        return

    # Highlight current sentence
    start, end = positions[index]
    text_box.tag_remove("sentence_highlight", "1.0", tk.END)
    text_box.tag_add("sentence_highlight", start, end)
    text_box.tag_configure("sentence_highlight", background="#FFFACD")
    text_box.see(start)

    # Update emoji
    bbox = text_box.bbox(start)
    margin_canvas.delete("ball")
    if bbox:
        x, y, w, h = bbox
        margin_canvas.create_text(15, y + h // 2, text="🟡", font=("Arial", 12), tags="ball")

    # Schedule next sentence
    if highlight_timer_id:
        root.after_cancel(highlight_timer_id)
    highlight_timer_id = root.after(intervals[index], lambda: highlight_sentences(positions, intervals, index + 1))




def generate_gtts_audio(text):
    try:
        start_time = time.time()
        tts = gTTS(text=text, lang="en", slow=False)
        duration = time.time() - start_time

        if duration > 6:
            print(f"⚠️ gTTS took {duration:.2f}s — may feel like a freeze.")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
            tts.save(tmp.name)
            return tmp.name

    except Exception as e:
        print("gTTS generation failed:", e)
        return None


def get_hybrid_intervals(sentences, total_audio_duration):
    # Heuristics:
    # - Base time per sentence (to account for natural pauses)
    # - Word count scaling (speech pace)
    # - Slight bias for very short sentences

    base_time_ms = 400  # give all sentences a starting buffer
    words = [len(s.split()) for s in sentences]
    total_words = sum(words)
    remaining_duration_ms = int(total_audio_duration * 1000) - (len(sentences) * base_time_ms)

    # Sanity check
    if remaining_duration_ms < 0:
        base_time_ms = int((total_audio_duration * 1000) / len(sentences))
        remaining_duration_ms = 0

    intervals = []
    for count in words:
        # Scale remaining time proportionally
        word_proportion = count / total_words if total_words else 1
        word_time = int(remaining_duration_ms * word_proportion)
        intervals.append(base_time_ms + word_time)

    return intervals



def show_loading():
    global loading_label
    if not loading_label:
        loading_label = tk.Label(root, text="🎤 Preparing voice...", font=(CHOSEN_FONT, 12), fg="#555", bg="#F0F8FF")
    loading_label.place(relx=0.5, rely=0.9, anchor="center")
    loading_label.update_idletasks()






def generate_and_continue(text, full_text, offset):
    global loading_label, loading_animating

    try:
        # Step 1: Generate audio with gTTS (this can block for 2–4 seconds)
        audio_path = get_gtts_audio(text)
        if not audio_path:
            hide_loading_spinner()
            fallback_offline_tts(text)
            return

        # Step 2: Parse and prep sentence data while we're waiting
        sentences = split_into_sentences(text)
        sentence_map = get_sentence_positions(sentences, full_text, offset=offset)
        audio_duration = MP3(audio_path).info.length
        intervals = get_hybrid_intervals(sentences, audio_duration)

        # Step 3: Hide spinner after prep is done
        hide_loading_spinner()

        # Step 4: Load and delay-start playback just a bit
        pygame.mixer.music.load(audio_path)
        root.after(200, pygame.mixer.music.play)  # ⏳ gives a smoother handoff

        # Step 5: Start highlighting & cleanup afterward
        root.after(300, lambda: wait_for_playback_and_start(sentence_map, intervals))
        root.after(int(audio_duration * 1000) + 500, lambda *_: cleanup_after_playback(audio_path))

        disable_replay()

    except Exception as e:
        hide_loading_spinner()
        print(f"❌ Error in generate_and_continue: {e}")
        fallback_offline_tts(text)




def hide_loading_spinner():
    global loading_animating, loading_label
    loading_animating = False
    if loading_label:
        loading_label.destroy()
        loading_label = None


def cancel_highlight():
    global highlight_timer_id
    if highlight_timer_id is not None:
        root.after_cancel(highlight_timer_id)
        highlight_timer_id = None
    text_box.tag_remove("sentence_highlight", "1.0", tk.END)
    margin_canvas.delete("ball")




    

def on_closing():
    pygame.mixer.music.stop()
    root.destroy()


#*******THIS IS THE END OF THE LISTEN FUNCTION**********


#========================================================


#*******THIS IS THE START OF THE WRITE FUNCTION**********


# Suppress system audio logs BEFORE importing sound libraries
def suppress_stderr():
    sys.stderr.flush()
    return open(os.devnull, 'w')


def restore_stderr(original_stderr):
    sys.stderr.close()
    sys.stderr = original_stderr

original_stderr = sys.stderr
sys.stderr = suppress_stderr()

# Now import sound-related modules
from gtts import gTTS
import pyttsx3
import pygame
import speech_recognition as sr
from mutagen.mp3 import MP3
from threading import Timer
import tkinter.font as tkfont




restore_stderr(original_stderr)

typing_buffer = {"word": None}
typing_timer = {"ref": None}
last_tts_timer = None

def start_voice_input():
    recognizer = sr.Recognizer()

    if not hasattr(start_voice_input, "status_label"):
        start_voice_input.status_label = tk.Label(write_tab, font=(CHOSEN_FONT, 11), fg="#444", bg="#FFFDF7")

    status_label = start_voice_input.status_label
    status_label.pack_forget()  # Make sure it's fresh
    status_label.config(text="🎙️ Listening...")
    status_label.pack(pady=10)

    write_tab.update_idletasks()

    try:
        with sr.Microphone() as source:
            recognizer.adjust_for_ambient_noise(source, duration=0.5)
            time.sleep(0.6)  # prevent clipping
            audio = recognizer.listen(source, timeout=5)

        status_label.config(text="🧠 Thinking...")
        write_tab.update_idletasks()

        text = recognizer.recognize_google(audio)
        write_text.insert(tk.END, text + " ")

        status_label.config(text="✅ Got it: “" + text.split()[0] + "…”")  # Optional: preview first word
        write_tab.after(2500, status_label.pack_forget)  # fade label away after a few seconds

    except sr.WaitTimeoutError:
        status_label.config(text="⏱️ No sound detected. Try again?")
        write_tab.after(2500, status_label.pack_forget)

    except sr.UnknownValueError:
        status_label.config(text="🤔 I couldn’t quite hear that.")
        write_tab.after(2500, status_label.pack_forget)

    except sr.RequestError as e:
        status_label.config(text="🚫 STT error. Try again later.")
        write_tab.after(2500, status_label.pack_forget)




def toggle_dyslexic():
    global dyslexic_mode
    dyslexic_mode = not dyslexic_mode

    try:
        tkfont.nametofont("OpenDyslexic")  # check if it exists
        chosen_font = "OpenDyslexic"
    except Exception:
        chosen_font = "Comic Sans MS"
        print("⚠️ 'OpenDyslexic' not found. Using fallback font.")

    font_name = chosen_font if dyslexic_mode else CHOSEN_FONT
    write_text.config(font=(font_name, 18))







tts_queue = []

def tts_worker():
    print("🌀 TTS worker started")
    while True:
        if tts_queue:
            text = tts_queue.pop(0)
            print(f"🎤 TTS saying: {text}")
            try:
                engine = pyttsx3.init()
                engine.say(text)
                engine.runAndWait()
            except Exception as e:
                print("TTS Error:", e)
        time.sleep(0.1)

# ✅ Prevent duplicate thread starts
if not tts_thread_started:
    threading.Thread(target=tts_worker, daemon=True).start()
    tts_thread_started = True




def speak_word_from_write_tab(text):
    print(f"🛠 Typing debounced word: {text}")
    
    # Cancel previous timer if typing continues
    if typing_timer["ref"]:
        typing_timer["ref"].cancel()

    def delayed_speak():
        if text:
            speak_with_gtts(text)  # ✅ Use working voice path

    t = Timer(0.3, delayed_speak)
    t.daemon = True
    t.start()
    typing_timer["ref"] = t





def color_new_char(event):
    index = write_text.index("insert-1c")
    write_text.tag_add("new_char", index, "insert")
    write_text.tag_config("new_char", foreground="#E91E63")  # pink or blue tone
    write_text.after(250, lambda: write_text.tag_remove("new_char", "1.0", tk.END))





def sound_out_word(event):
    if event.keysym == "space":
        current_text = write_text.get("1.0", "insert").split()
        if current_text:
            last_word = current_text[-1]
            speak_word_from_write_tab(last_word)





def get_clicked_word(event):
    index = write_text.index(f"@{event.x},{event.y}")
    line_text = write_text.get(f"{index} linestart", f"{index} lineend")

    # Figure out clicked char position within that line
    column = int(index.split(".")[1])
    words = list(re.finditer(r"\b\w+\b", line_text))

    for match in words:
        if match.start() <= column <= match.end():
            return match.group()
    return None



def speak_word(event):
    word = get_clicked_word(event)
    if word:
        speak_word_from_write_tab(word)


def highlight_word(word):
    start = write_text.search(word, "1.0", tk.END)
    if start:
        end = f"{start}+{len(word)}c"
        write_text.tag_add("speak_highlight", start, end)
        write_text.tag_config("speak_highlight", background="#FFF176")
        root.after(600, lambda: write_text.tag_remove("speak_highlight", "1.0", tk.END))


def test_typing_voice():
    speak_word_from_write_tab("testing")





# Initialize mixer once at the start
pygame.mixer.init()

def speak_with_gtts(text):
    try:
        tts = gTTS(text=text, lang='en')
        temp_file = "mia_temp.mp3"
        tts.save(temp_file)

        pygame.mixer.music.load(temp_file)
        pygame.mixer.music.play()

        while pygame.mixer.music.get_busy():
            pygame.time.Clock().tick(10)  # Avoid CPU overuse

        if os.path.exists(temp_file):
            os.remove(temp_file)

    except Exception as e:
        print(f"⚠️ TTS failed: {e}")
        # Optional: show a messagebox or log error
        # messagebox.showwarning("TTS Error", "Unable to play audio.")



#*******START OF LISTEN AND TYPE CHALLENGE**********


def load_challenge_wordlist(file_path="challenge_words.txt"):
    try:
        with open(file_path) as f:
            words = [line.strip().lower() for line in f if line.strip()]
        return words
    except FileNotFoundError:
        messagebox.showerror("Missing File", f"Could not find wordlist: {file_path}")
        return []

challenge_wordlist = load_challenge_wordlist()
challenge_index = 0
challenge_score = {"correct": 0, "total": 0}



def upload_challenge_wordlist():
    file_path = filedialog.askopenfilename(
        title="Select Challenge Wordlist",
        filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
    )
    if not file_path:
        return

    try:
        with open(file_path) as f:
            words = [line.strip().lower() for line in f if line.strip()]
        if not words:
            messagebox.showwarning("Empty File", "No words found in the selected file.")
            return

        # Reset global wordlist and index
        global challenge_wordlist, challenge_index, challenge_score
        challenge_wordlist = words
        challenge_index = 0
        challenge_score = {"correct": 0, "total": 0}

        messagebox.showinfo("✅ Wordlist Loaded", f"{len(words)} words loaded!\nReady to start a new challenge.")
        feedback_label.config(text="📂 New wordlist uploaded!", fg="#333")
    except Exception as e:
        messagebox.showerror("❌ Error", f"Could not load file:\n{e}")






def start_listen_and_type_challenge():
    global challenge_index
    if challenge_index >= len(challenge_wordlist):
        messagebox.showinfo("🏁 Done!", "You've finished all words!")
        return

    word = challenge_wordlist[challenge_index]
    write_text.delete("1.0", tk.END)
    feedback_label.config(text=f"🎧 Listen and type: {word}", fg="#333")

    # Use voice prompt — you can swap this for your own phonogram playback!
    speak_with_gtts(word)




def check_challenge_answer():
    global challenge_index
    typed = write_text.get("1.0", tk.END).strip().lower()
    target = challenge_wordlist[challenge_index]

    challenge_score["total"] += 1
    write_text.delete("1.0", tk.END)

    if typed == target:
        challenge_score["correct"] += 1
        feedback_label.config(text="✅ Great job!", fg="#2E7D32")
        challenge_index += 1
        write_text.insert(tk.END, target)
    else:
        # 🎤 Say correction
        speak_with_gtts(f"You typed {typed}. The correct spelling is {target}.")

        # 📝 Show and highlight differences
        write_text.insert(tk.END, "")
        for i in range(max(len(typed), len(target))):
            typed_char = typed[i] if i < len(typed) else ""
            target_char = target[i] if i < len(target) else ""

            if typed_char == target_char:
                write_text.insert(tk.END, target_char, "correct")
            else:
                write_text.insert(tk.END, target_char, "error")

        write_text.tag_config("correct", foreground="#388E3C")  # green
        write_text.tag_config("error", foreground="#D32F2F")    # red

        feedback_label.config(
            text=f"❌ Try again.\nCorrect spelling: {target}", fg="#E53935")




def show_progress():
    correct = challenge_score["correct"]
    total = challenge_score["total"]
    messagebox.showinfo("📈 Progress", f"You've answered {correct} out of {total} correctly!")





#*******THIS IS THE END OF THE WRITE FUNCTION  (on_key_release IS PART OF THIS FUNCTION)**********



#========================================================


#*******THIS IS THE START OF THE GAME FUNCTION**********

import random
from tkinter import filedialog
import string
import os
from docx import Document
import openpyxl
import json
import tkinter.simpledialog as simpledialog
import hashlib

result_label = None
spell_frame = None
back_btn = None
session_results = []  # Tracks "correct" / "incorrect" answers this round


WORDLIST_DIR = "wordlists"
os.makedirs(WORDLIST_DIR, exist_ok=True)


CONFIG_FILE = "litbuddy_config.json"

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    return {"active_list": "default.json"}

def save_config(data):
    with open(CONFIG_FILE, "w") as f:
        json.dump(data, f)

config = load_config()


def get_wordlist_path(filename):
    return os.path.join(WORDLIST_DIR, filename)


def load_spell_words():
    path = config["active_list"]
    if not os.path.isabs(path):
        path = os.path.join(WORDLIST_DIR, path)

    if os.path.exists(path):
        with open(path, "r") as f:
            return json.load(f)
    else:
        return []


spell_words = load_spell_words()

def confuse(word):
    if len(word) < 2:
        return word  # too short to mess with

    index = random.randint(0, len(word) - 1)
    letter = random.choice(string.ascii_lowercase.replace(word[index], ""))
    return word[:index] + letter + word[index+1:]




def extract_words_from_file(file_path):
    ext = os.path.splitext(file_path)[-1].lower()

    if ext in [".txt", ".csv"]:
        with open(file_path, "r") as f:
            return f.read()

    elif ext == ".docx":
        doc = Document(file_path)
        return "\n".join(p.text for p in doc.paragraphs)

    elif ext == ".xlsx":
        wb = openpyxl.load_workbook(file_path)
        sheet = wb.active
        words = []
        for row in sheet.iter_rows(values_only=True):
            for cell in row:
                if cell:
                    words.append(str(cell))
        return "\n".join(words)

    else:
        return ""  


def upload_word_file():
    global spell_words
    file_path = filedialog.askopenfilename(filetypes=[
        ("Text files", "*.txt"),
        ("CSV files", "*.csv"),
        ("Word documents", "*.docx"),
        ("Excel files", "*.xlsx")
    ])
    if not file_path:
        return

    content = extract_words_from_file(file_path)
    if not content.strip():
        messagebox.showerror("⚠️ No Words", "Could not extract any words from the selected file.")
        return

    new_words = convert_word_list(content)
    print("WORDS ADDED:", new_words)  # 🔍 Debug print to check

    spell_words.extend(new_words)
    save_spell_words()

    speak_with_gtts("Your new words were added!")
    messagebox.showinfo("✅ Upload Complete", f"{len(new_words)} new words added.")





def convert_word_list(raw_text):
    words = [line.strip().lower() for line in raw_text.splitlines() if line.strip()]
    if not words:
        print("❌ No valid words found in upload.")
        return []

    converted = []
    for w in words:
        syllables = dic.inserted(w).split('-')
        converted.append({
            "word": w,
            "choices": [w, confuse(w), confuse(w)],
            "syllables": syllables  # <-- this is key
        })
    return converted



def save_spell_words():
    path = config["active_list"]
    if not os.path.isabs(path):
        path = os.path.join(WORDLIST_DIR, path)

    with open(path, "w") as f:
        json.dump(spell_words, f, indent=2)
    print(f"✅ Saved {len(spell_words)} words to {path}")




def open_wordlist_manager():
    manager = tk.Toplevel()
    manager.title("📁 Manage Word Lists")

    tk.Label(manager, text="Select a list:").pack(pady=(10, 2))

    listbox = tk.Listbox(manager, height=10, width=30)
    listbox.pack()

    # Load all .json files in folder
    lists = [f for f in os.listdir(WORDLIST_DIR) if f.endswith(".json")]
    for f in lists:
        listbox.insert(tk.END, f)

    
    def set_active_list():
        selection = listbox.curselection()
        if selection:
            selected_file = listbox.get(selection[0])
            full_path = os.path.join(WORDLIST_DIR, selected_file)

            global config, spell_words
            config["active_list"] = selected_file  # ✅ filename only

            # ✅ Save active list into user's profile (if applicable)
            profiles_data = load_profiles()
            user = profiles_data.get("current_user")
            if user and user in profiles_data["profiles"]:
                profiles_data["profiles"][user]["active_list"] = selected_file
                save_profiles(profiles_data)

            save_config(config)

            try:
                with open(full_path) as f:
                    spell_words = json.load(f)

                # ✅ Safely update the label if it exists
                try:
                    if active_list_label and active_list_label.winfo_exists():
                        active_list_label.config(text=f"📂 Current List: {selected_file}")
                except Exception as e:
                    print("Label update skipped:", e)

                speak_with_gtts("Now using " + selected_file)
                messagebox.showinfo("✅ List Switched", f"Now using '{selected_file}'")
                manager.destroy()

            except Exception as e:
                messagebox.showerror("❌ Error", f"Could not load selected list:\n{e}")
                return


           

    def delete_selected():
        selection = listbox.curselection()
        if selection:
            selected = listbox.get(selection[0])
            path = get_wordlist_path(selected)
            if messagebox.askyesno("Delete?", f"Delete {selected}?"):
                os.remove(path)
                listbox.delete(selection[0])




    def create_new_list():
        file_path = filedialog.askopenfilename(title="Add Wordlist from File", filetypes=[
            ("All Supported", "*.txt *.csv *.docx *.xlsx"),
            ("Text Files", "*.txt"),
            ("CSV Files", "*.csv"),
            ("Word Docs", "*.docx"),
            ("Excel Files", "*.xlsx")
        ])
        if not file_path:
            return

        # Extract raw word content
        raw_text = extract_words_from_file(file_path)
        if not raw_text.strip():
            messagebox.showerror("⚠️ No Words", "No words were found in the selected file.")
            return

        # Convert to spell game format
        new_words = convert_word_list(raw_text)
        if not new_words:
            messagebox.showwarning("⚠️ Skip", "No valid words to convert.")
            return

        # Save to wordlists/
        base = os.path.splitext(os.path.basename(file_path))[0]
        output_file = os.path.join(WORDLIST_DIR, f"{base}_game_ready.json")
        with open(output_file, "w") as f:
            json.dump(new_words, f, indent=2)

        listbox.insert(tk.END, os.path.basename(output_file))
        messagebox.showinfo("✅ Wordlist Ready", f"{len(new_words)} words saved to:\n{os.path.basename(output_file)}")


    tk.Button(manager, text="📄 New List", command=create_new_list).pack(pady=5)
    tk.Button(manager, text="➕ Add Wordlist from File...", command=create_new_list).pack(pady=5)
    tk.Button(manager, text="🗂️ Use Selected", command=set_active_list).pack()
    tk.Button(manager, text="🗑️ Delete Selected", command=delete_selected).pack(pady=(5, 10))




def hash_(password):
    return hashlib.sha256(password.encode()).hexdigest()


def load_admin_config():
    with open("admin_config.json", "r") as f:
        return json.load(f)

def check_password():
    try:
        with open("admin_config.json", "r") as f:
            admin_data = json.load(f)
    except FileNotFoundError:
        messagebox.showerror("Config Missing", "admin_config.json not found.")
        return False

    # If no password is set, prompt to create one
    if not admin_data.get("password_hash"):
        messagebox.showinfo("🔐 Setup", "No admin password is set. Let's create one.")
        new = simpledialog.askstring("🔑 New Password", "Enter new password:", show="*")
        if not new:
            return False
        confirm = simpledialog.askstring("🔁 Confirm Password", "Re-enter new password:", show="*")
        if new != confirm:
            messagebox.showerror("❗ Mismatch", "Passwords do not match.")
            return False

        new_hash = hashlib.sha256(new.encode()).hexdigest()
        with open("admin_config.json", "w") as f:
            json.dump({"password_hash": new_hash}, f)
        messagebox.showinfo("✅ Set", "Admin password created successfully!")
        return True

    # Otherwise, prompt for existing password
    entered = simpledialog.askstring("🔐 Admin Access", "Enter password:", show="*")
    if not entered:
        return False

    entered_hash = hashlib.sha256(entered.encode()).hexdigest()
    return entered_hash == admin_data.get("password_hash")




def password_protected_manager():
    if check_password():
        open_wordlist_manager()
    else:
        messagebox.showwarning("Access Denied", "Incorrect password.")


def change_admin_password():
    # Ask for current password
    current = simpledialog.askstring("🔐 Current Password", "Enter current password:", show="*")
    if not current:
        return

    admin_data = load_admin_config()
    current_hash = hashlib.sha256(current.encode()).hexdigest()
    if current_hash != admin_data.get("password_hash"):
        messagebox.showerror("❌ Incorrect", "Current password is incorrect.")
        return

    # Ask for new password (twice to confirm)
    new = simpledialog.askstring("🔑 New Password", "Enter new password:", show="*")
    if not new:
        return
    confirm = simpledialog.askstring("🔁 Confirm Password", "Re-enter new password:", show="*")
    if new != confirm:
        messagebox.showerror("❗ Mismatch", "Passwords do not match.")
        return

    # Save the new hash
    new_hash = hashlib.sha256(new.encode()).hexdigest()
    with open("admin_config.json", "w") as f:
        json.dump({"password_hash": new_hash}, f)
    messagebox.showinfo("✅ Updated", "Password changed successfully!")




#******🔊 Audio Challenge*******

import pyphen
dic = pyphen.Pyphen(lang='en')

def get_phonetic_hint(word):
    return dic.inserted(word).replace('-', '-')


def start_audio_game():
    global result_label, back_btn

    clear_play_area()
    menu_frame.pack_forget()
    game_label.config(text="🔊 Audio Challenge!")

    # ✅ Result label
    result_label = tk.Label(play_tab, text="", font=(CHOSEN_FONT, 14), bg="#FFF9F2")
    result_label.pack(pady=10)

    # 🔙 Back to Menu button
    
    back_btn = tk.Button(play_tab, text="🔙 Back to Menu", font=(CHOSEN_FONT, 12), bg="#FFDDDD", command=show_menu, state="disabled")
    back_btn.pack(pady=(0, 10))


    load_audio_word()


def load_audio_word():
    global spell_frame, result_label, back_btn, spell_words

    if result_label:
        result_label.config(text="")

    if not spell_words:
        messagebox.showerror("🚫 No Words Found", "The current word list is empty.\nPlease upload words or select a different list.")
        if back_btn and back_btn.winfo_exists():
            back_btn.config(state="normal")
        return

    if spell_frame:
        spell_frame.destroy()

    spell_frame = tk.Frame(play_tab, bg="#F9F9F2")
    spell_frame.pack()

    entry = random.choice(spell_words)
    correct = entry["word"]
    choices = random.sample(entry["choices"], len(entry["choices"]))

    profiles = load_profiles()
    user = profiles.get("current_user", "Player")

    # 🔒 Disable Back button while speaking
    if back_btn and back_btn.winfo_exists():
        back_btn.config(state="disabled")

    # 🗣️ Speak the word and then re-enable the back button
    def speak_and_enable():
        speak_with_gtts(f"The word is: {correct}")
        if back_btn and back_btn.winfo_exists():
            back_btn.config(state="normal")

    play_tab.after(300, speak_and_enable)

    # 🔁 Repeat word button
    repeat_btn = tk.Button(spell_frame, text="🔁 Repeat Word", font=(CHOSEN_FONT, 11), bg="#E6E6FA",
                           command=lambda: speak_with_gtts(f"The word is: {correct}"))
    repeat_btn.pack(pady=(0, 10))

    # ➕ Create answer buttons
    for choice in choices:
        btn = tk.Button(spell_frame, text=choice, font=(CHOSEN_FONT, 14), width=12,
                        command=lambda c=choice: check_audio_answer(c, correct))
        btn.pack(pady=6)


AUDIO_MILESTONES = {
    "words_attempted": [5, 15, 30],
    "correct_answers": [5, 10, 20]
}

def check_audiochallenge_milestone(user):
    profiles = load_profiles()
    profile = profiles["profiles"].get(user, {})
    progress = profile.get("progress", {})
    ac_milestones = profile.setdefault("audio_milestones", {
        "words_attempted": [],
        "correct_answers": []
    })

    attempted = progress.get("games_played", 0)
    correct = progress.get("correct", 0)
    triggered = False

    for target in AUDIO_MILESTONES["words_attempted"]:
        if attempted >= target and target not in ac_milestones["words_attempted"]:
            show_congrats_banner(f"🔊 You've completed {target} audio challenges!")
            ac_milestones["words_attempted"].append(target)
            triggered = True
            break

    for target in AUDIO_MILESTONES["correct_answers"]:
        if correct >= target and target not in ac_milestones["correct_answers"]:
            show_congrats_banner(f"🎯 {correct} correct answers in Audio Challenge!")
            ac_milestones["correct_answers"].append(target)
            triggered = True
            break

    if triggered:
        save_profiles(profiles)



def check_audio_answer(selected, correct):
    global result_label, back_btn

    user = load_profiles().get("current_user")

    if selected == correct:
        result_label.config(text="🎉 You got it!", fg="#2e7d32")
        speak_with_gtts("Correct! Great listening!")
        update_progress(user, correct=True)
    else:
        result_label.config(text="😬 Try again!", fg="#c62828")
        speak_with_gtts("Oops! Not quite.")

    save_profiles(load_profiles())
    check_audiochallenge_milestone(user)

    back_btn.config(state="normal")  # ✅ Unlock the menu button
    after_ids["audio_game"] = play_tab.after(2000, load_audio_word)





#******Start of Syllable Snap! Challenge*******

active_list_label = None




def start_syllable_snap():
    global result_label, back_btn, active_list_label

    clear_play_area()
    menu_frame.pack_forget()
    game_label.config(text="🧩 Syllable Snap!")

    # 📂 Show active word list
    active_list_label = tk.Label(play_tab, text=f"📂 Current List: {config['active_list']}", font=(CHOSEN_FONT, 12), bg="#FFF9F2", fg="#555")
    active_list_label.pack(pady=(5, 0))

    result_label = tk.Label(play_tab, text="", font=(CHOSEN_FONT, 14), bg="#FFF9F2")
    result_label.pack(pady=10)

    back_btn = tk.Button(play_tab, text="🔙 Back to Menu", font=(CHOSEN_FONT, 12), bg="#FFDDDD",
                         command=show_menu, state="disabled")
    back_btn.pack(pady=(0, 10))

    load_syllable_snap_round()





def load_syllable_snap_round():
    global spell_frame, result_label, back_btn, spell_words

    # 🧹 Clear any previous frame
    if spell_frame:
        spell_frame.destroy()

    # 🎯 Ensure result_label exists and is reset
    if not result_label or not result_label.winfo_exists():
        result_label = tk.Label(play_tab, font=(CHOSEN_FONT, 14), bg="#F9F9FF")
        result_label.pack(pady=10)
    else:
        result_label.config(text="")

    # ✅ Always enable the Back to Menu button
    if back_btn and back_btn.winfo_exists():
        back_btn.config(state="normal")

    # 🔍 Filter for valid syllable words
    valid_syllable_words = [w for w in spell_words if "syllables" in w and len(w["syllables"]) > 1]

    if not valid_syllable_words:
        result_label.config(text="⚠️ No words with syllables found. Upload or convert a new list.", fg="#E53935")
        return

    # 🧠 Select a word and its syllables
    entry = random.choice(valid_syllable_words)
    word = entry["word"]
    syllables = entry["syllables"]

    if not word or not syllables:
        result_label.config(text="⚠️ Missing syllables for this word.", fg="#E53935")
        return

    shuffled = syllables[:]
    random.shuffle(shuffled)

    # 🧩 Create new frame for buttons
    spell_frame = tk.Frame(play_tab, bg="#FFF9F2")
    spell_frame.pack()

    result_label.config(text=f"🔊 Segment the word: {word}", fg="#333")

    selected = []

    def handle_syllable_click(syl):
        selected.append(syl)

        if len(selected) == len(syllables):
            if selected == syllables:
                result_label.config(text="🎉 Correct!", fg="#2e7d32")
                speak_with_gtts(f"Well done! {word}")
            else:
                result_label.config(text="😬 Not quite. Try again!", fg="#c62828")
                speak_with_gtts("Hmm, not the right order.")

            #play_tab.after(2000, load_syllable_snap_round)
            
            after_ids["syllable_snap"] = play_tab.after(2000, load_syllable_snap_round)



    for syl in shuffled:
        btn = tk.Button(spell_frame, text=syl, font=(CHOSEN_FONT, 14), width=10,      command=lambda s=syl: handle_syllable_click(s))
        btn.pack(pady=5)





#******End of Syllable Snap! Challenge*******

#============================================


#******Start of Sound It Out!*******

import pygame
import time
import subprocess




def threaded_speak(text):
    threading.Thread(target=lambda: speak_chunk(text)).start()


def synthesize_phoneme_audio(text, output_dir="phonemes_audio"):
    import hashlib
    os.makedirs(output_dir, exist_ok=True)

    # Use hash to avoid illegal filename characters
    safe_name = hashlib.md5(text.encode()).hexdigest()
    filename = os.path.join(output_dir, f"{safe_name}.wav")

    if not os.path.exists(filename):
        try:
            subprocess.run(["espeak-ng", "-v", "en", "-w", filename, text])
        except Exception as e:
            print(f"Error synthesizing '{text}': {e}")
    return filename


def speak_chunk(text):
    audio_path = synthesize_phoneme_audio(text)

    pygame.mixer.init()
    if os.path.isfile(audio_path):
        pygame.mixer.music.load(audio_path)
        pygame.mixer.music.play()
        while pygame.mixer.music.get_busy():
            time.sleep(0.1)
    else:
        print(f"⚠️ Missing audio for {text}")





def load_phonics_map():
    try:
        with open(PHONICS_MAP_PATH) as f:
            raw = json.load(f)
            # Convert list of dicts to a single dict if needed
            if isinstance(raw, list):
                return {k: v for d in raw for k, v in d.items()}
            elif isinstance(raw, dict):
                return raw
            else:
                print("⚠️ Unexpected phonics map format.")
                return {}
    except Exception as e:
        print(f"Error loading phonics map: {e}")
        return {}

phonics_audio_map = load_phonics_map()


def choose_phonics_list():
    path = filedialog.askopenfilename(
        title="Select Phonics Wordlist",
        initialdir=WORDLIST_DIR,
        filetypes=[("JSON files", "*.json")]
    )
    if path:
        config["phonics_list"] = os.path.basename(path)
        save_config(config)
        messagebox.showinfo("✅ Ready", f"{os.path.basename(path)} will now be used for Sound It Out.")


def start_phonics_game():
    global result_label, back_btn

    clear_play_area()
    menu_frame.pack_forget()
    game_label.config(text="🔤 Sound It Out!")

    result_label = tk.Label(play_tab, text="", font=(CHOSEN_FONT, 14), bg="#FFF9F2") 
    result_label.pack(pady=10)

    back_btn = tk.Button(play_tab, text="🔙 Back to Menu", font=(CHOSEN_FONT, 12), bg="#FFDDDD", command=show_menu, state="disabled")
    back_btn.pack(pady=(0, 10))

    selected_file = config.get("phonics_list")
    if not selected_file or not os.path.isfile(get_wordlist_path(selected_file)):
        messagebox.showwarning("📂 No Wordlist Found", "Please select a wordlist to begin.")
        choose_phonics_list()
        selected_file = config.get("phonics_list")

    phonics_words = []
    try:
        with open(get_wordlist_path(selected_file)) as f:
            phonics_words = json.load(f)
    except Exception as e:
        result_label.config(text=f"⚠️ Failed to load phonics list: {e}", fg="#E53935")
        back_btn.config(state="normal")
        return

    if not phonics_words:
        result_label.config(text="⚠️ Wordlist is empty or invalid.", fg="#E53935")
        back_btn.config(state="normal")
        return

    load_phonics_game_round(phonics_words)




def enrich_wordlist_with_phonemes(wordlist):
    for entry in wordlist:
        if "phonics" not in entry or not entry["phonics"]:
            phones = pronouncing.phones_for_word(entry["word"])
            if phones:
                phonemes = phones[0].split()
                entry["phonics"] = phonemes_to_syllables(phonemes)
                entry["ipa"] = phonemes_to_ipa(phonemes)
            else:
                entry["phonics"] = [entry["word"]]
                entry["ipa"] = ""
    return wordlist

#load_phonics_game_round(phonics_words)


def get_phonemes_espeak(word):
    try:
        result = subprocess.run(
            ["espeak-ng", "-q", "--ipa=3", word],
            capture_output=True,
            text=True
        )
        phonemes = result.stdout.strip()
        return phonemes.split()  # Split into phoneme chunks
    except Exception as e:
        print(f"Error generating phonemes: {e}")
        return []


def play_sound_with_feedback(btn, grapheme):
    original_color = btn["bg"]
    btn.config(bg="#FFF176")  # highlight
    threaded_speak(grapheme)
    play_tab.after(600, lambda: btn.config(bg=original_color))

def upload_and_convert_wordlist():
            path = filedialog.askopenfilename(
                title="Upload Wordlist",
                filetypes=[("Text files", "*.txt"), ("CSV files", "*.csv"), ("JSON files", "*.json")]
            )
            if not path:
                return

            try:
                ext = os.path.splitext(path)[1].lower()
                words = []

                # 📥 Load raw words from file
                if ext == ".txt":
                    with open(path) as f:
                        words = [line.strip() for line in f if line.strip()]
                elif ext == ".csv":
                    import csv
                    with open(path) as f:
                        reader = csv.reader(f)
                        words = [row[0].strip() for row in reader if row]
                elif ext == ".json":
                    with open(path) as f:
                        raw = json.load(f)
                        if isinstance(raw, list):
                            words = [entry["word"] for entry in raw if "word" in entry]
                        elif isinstance(raw, dict):
                            words = list(raw.keys())

                # 🧠 Enrich with syllables + IPA
                from ipa_converter import generate_data
                data = generate_data(words)
                enriched = [{"word": w, "phonics": info["syllables"], "ipa": info["ipa"]} for w, info in data.items()]

                # 💾 Save to JSON
                save_path = os.path.join(WORDLIST_DIR, "uploaded_wordlist.json")
                with open(save_path, "w") as f:
                    json.dump(enriched, f, indent=2)

                # 🔧 Update config
                config["phonics_list"] = "uploaded_wordlist.json"
                save_config(config)
                messagebox.showinfo("✅ Wordlist Ready", "Your uploaded wordlist is now ready to use!")

            except Exception as e:
                messagebox.showerror("❌ Error", f"Failed to process wordlist: {e}")

def reload_phonics_wordlist():
    choose_phonics_list()
    selected_file = config.get("phonics_list")
    if not selected_file or not os.path.isfile(get_wordlist_path(selected_file)):
        messagebox.showerror("❌ Error", "No valid wordlist selected.")
        return

    try:
        with open(get_wordlist_path(selected_file)) as f:
            phonics_words = json.load(f)
        load_phonics_game_round(phonics_words)
    except Exception as e:
        messagebox.showerror("❌ Error", f"Failed to reload wordlist: {e}")





def load_phonics_game_round(phonics_words):
    global spell_frame, result_label, back_btn

    if spell_frame:
        spell_frame.destroy()

    spell_frame = tk.Frame(play_tab, bg="#FFF9F2")
    spell_frame.pack()

    if result_label:
        result_label.config(text="")

    valid_phonics_words = [w for w in phonics_words if "phonics" in w and len(w["phonics"]) > 1]
    if not valid_phonics_words:
        result_label.config(text="⚠️ No phonics-ready words found!", fg="#E53935")
        back_btn.config(state="normal")
        return

    entry = random.choice(valid_phonics_words)
    word = entry["word"]
    phonics = entry["phonics"]

    # IPA display
    if "ipa" in entry:
        ipa_label = tk.Label(spell_frame, text=f"IPA: /{entry['ipa']}/", font=(CHOSEN_FONT, 12), bg="#FFF9F2", fg="#616161")
        ipa_label.pack(pady=(0, 5))

    # Speak intro and phonics
    speak_chunk("let's sound it out")
    for p in phonics:
        speak_chunk(p)

    result_label.config(text=f"🔊 Sound it out: {word}", fg="#333")

    def play_sound(grapheme):
        speak_chunk(grapheme)

    for ph in phonics:
        btn = tk.Button(spell_frame, text=ph, font=(CHOSEN_FONT, 14), width=6)
        btn.config(command=lambda b=btn, s=ph: play_sound_with_feedback(b, s))
        btn.pack(pady=5)

    # Blend button
    blend_btn = tk.Button(spell_frame, text="🧩 Blend Sounds Together", font=(CHOSEN_FONT, 12), bg="#FFFDE7", command=lambda: blend_sounds_together(phonics, word))
    blend_btn.pack(pady=5)

    # Feedback after delay
    play_tab.after(2000, lambda: result_label.config(text="🎉 Great job! Want to try a new word?", fg="#2e7d32"))

    # Say full word
    say_word_btn = tk.Button(spell_frame, text="▶️ Say Full Word", font=(CHOSEN_FONT, 12), bg="#E3F2FD", command=lambda: speak_chunk(word))
    say_word_btn.pack(pady=5)

    # Next word
    next_word_btn = tk.Button(spell_frame, text="🔁 Next Word", font=(CHOSEN_FONT, 12), bg="#E8F5E9", command=lambda: load_phonics_game_round(phonics_words))
    next_word_btn.pack(pady=5)

    # Upload wordlist
    upload_btn = tk.Button(spell_frame, text="📂 Upload Wordlist", font=(CHOSEN_FONT, 12), bg="#D1C4E9", command=upload_and_convert_wordlist)
    upload_btn.pack(pady=5)

    # Reload wordlist
    reload_btn = tk.Button(spell_frame, text="🔄 Reload Wordlist", font=(CHOSEN_FONT, 12), bg="#FFF3E0", command=reload_phonics_wordlist)
    reload_btn.pack(pady=5)

    back_btn.config(state="normal")




    





def blend_sounds_together(phonics, word):
    def play_sequence():
        for p in phonics:
            speak_chunk(p)
            time.sleep(0.5)
        speak_chunk(f"now say it together: {word}")
    threading.Thread(target=play_sequence).start()





#******End of Sound It Out!*******




#*******THIS IS THE END OF THE GAME FUNCTION**********


#=======================================================

#*******THIS IS THE START OF THE PROGRESS FUNCTION**********


def update_progress(user, correct=False, incorrect=False, game_completed=False):
    profiles_data = load_profiles()

    if user not in profiles_data["profiles"]:
        return  # No such user

    # Get actual user profile reference
    user_profile = profiles_data["profiles"]
    check_milestones(user)



def show_progress_ui(container):  
    for widget in container.winfo_children():
        widget.destroy()  # ✅ Clear previous content

    profiles_data = load_profiles()
    current_user = profiles_data.get("current_user")
    if not current_user:
        tk.Label(container, text="No profile selected.").pack()
        return

    progress = profiles_data["profiles"][current_user].get("progress", {})
    played = progress.get("games_played", 0)
    correct = progress.get("correct", 0)
    incorrect = progress.get("incorrect", 0)

    # ✅ Now populate cleanly
    tk.Label(container, text=f"🎓 {current_user}'s Progress", font=(CHOSEN_FONT, 20)).pack(pady=(10, 5))
    tk.Label(container, text=f"📝 Spelling Words Attempted: {played}", font=(CHOSEN_FONT, 12)).pack()
    tk.Label(container, text=f"✅ Correct Answers: {correct}", font=(CHOSEN_FONT, 12), fg="#2b8c34").pack()
    tk.Label(container, text=f"❌ Incorrect Answers: {incorrect}", font=(CHOSEN_FONT, 12), fg="#b22222").pack()    

    # 🔁 Refresh Button
    tk.Button(container, text="🔄 Refresh Progress", font=(CHOSEN_FONT, 10), bg="#e0e0e0", command=lambda: refresh_progress(container)).pack(pady=(10, 2))

    # ✏️ Reset Button (admin only)
    tk.Button(container, text="✏️ Reset Stats", font=(CHOSEN_FONT, 10), bg="#ffcaca", fg="#a00000", command=lambda: reset_user_progress(container)).pack()




def refresh_progress(container):
    for widget in container.winfo_children():
        widget.destroy()
    show_progress_ui(container)  # ✅ show updated stats here



def reset_user_progress(container):
    profiles_data = load_profiles()
    current_user = profiles_data.get("current_user")

    if not current_user:
        messagebox.showwarning("⚠️ No User", "No profile selected.")
        return

    # 🔐 Admin password check
    if not check_password():
        messagebox.showerror("Access Denied", "Incorrect admin password.")
        return

    confirm = messagebox.askyesno("Reset Stats", f"Are you sure you want to reset {current_user}'s stats?")
    if not confirm:
        return

    profiles_data["profiles"][current_user]["progress"] = {
        "games_played": 0,
        "correct": 0,
        "incorrect": 0
    }

    save_profiles(profiles_data)
    refresh_progress(container)
    messagebox.showinfo("✅ Stats Reset", f"{current_user}'s stats have been cleared.")


def show_congrats_banner(message):
    banner = tk.Toplevel()
    banner.title("🎉 Milestone Achieved!")
    banner.geometry("300x150")
    banner.configure(bg="#fffbcc")

    tk.Label(banner, text=message, font=(CHOSEN_FONT, 14), fg="#2b8c34", bg="#fffbcc").pack(expand=True, pady=30)
    tk.Button(banner, text="Yay!", command=banner.destroy, bg="#c1e1c1").pack(pady=10)

    # Optionally make it auto-close after a few seconds:
    banner.after(4000, banner.destroy)




MILESTONES = {
    "words_attempted": [10, 25, 50, 100],
    "correct_answers": [5, 20, 40, 75]
}


def check_milestones(user, container=None):
    data = load_profiles()
    profile = data["profiles"].get(user, {})
    progress = profile.get("progress", {})
    attempted = progress.get("games_played", 0)  # renamed in display only
    correct = progress.get("correct", 0)

    if attempted in MILESTONES["words_attempted"]:
        show_congrats_banner(f"🧩 You've attempted {attempted} words! Keep going!")

    if correct in MILESTONES["correct_answers"]:
        show_congrats_banner(f"✅ {correct} correct answers! Great job!")

def show_report_card():
        report = f"""
    📋 Report Card for {current_user}

    ➕ Add It Up:
      - Attempts: {addition_score['attempts']}
      - Correct: {addition_score['correct']}
      - Accuracy: {round((addition_score['correct'] / addition_score['attempts']) * 100, 1) if addition_score['attempts'] else 0}%

    🎯 Type the Word:
      - Attempts: {typing_score['attempts']}
      - Correct: {typing_score['correct']}
      - Accuracy: {round((typing_score['correct'] / typing_score['attempts']) * 100, 1) if typing_score['attempts'] else 0}%
    """

        report_window = tk.Toplevel()
        report_window.title("📋 Report Card")
        report_window.geometry("400x300")
        report_window.configure(bg="#FFF")

        tk.Label(report_window, text=report, font=(CHOSEN_FONT, 12), justify="left", bg="#FFF").pack(padx=20, pady=20)
        tk.Button(report_window, text="Close", command=report_window.destroy).pack(pady=10)

        tk.Button(addition_tab, text="📋 Show Report Card", font=(CHOSEN_FONT, 11), bg="#FFF9C4", command=show_report_card).pack(pady=5)







# === GUI Setup ===
root = tk.Tk()
root.title("📚 LitBuddy – Let's Learn Together!")
root.geometry("700x500")
root.configure(bg="#F0F8FF")
root.protocol("WM_DELETE_WINDOW", on_closing)

settings_visible = tk.BooleanVar(value=False)



notebook = ttk.Notebook(root)
notebook.pack(expand=True, fill="both", padx=10, pady=10)

#Listen Tab
listen_tab = tk.Frame(notebook, bg="#F0F8FF")
notebook.add(listen_tab, text="🔊 Listen")

# Write tab
write_tab = tk.Frame(notebook, bg="#FFFDF7")  # soft warm tone for writing
notebook.add(write_tab, text="✍️ Write")

# Play Tab
play_tab = tk.Frame(notebook, bg="#F9F9FF")
notebook.add(play_tab, text="🎮 Play")

# 📊 Progress Tab
progress_tab = tk.Frame(notebook, bg="#f7faff")
notebook.add(progress_tab, text="📊 My Progress")

# 🧠 Inject Progress Content
show_progress_ui(progress_tab)


# Teach Tab

teach_tab = tk.Frame(notebook, bg="#FAF3E0")  # warm background for educators
notebook.add(teach_tab, text="👩‍🏫 Teach")


#*******START of TEACH TAB**********
tk.Label(teach_tab, text="👩‍🏫 Teacher Controls", font=(CHOSEN_FONT, 16, "bold"), bg="#FAF3E0", fg="#333").pack(pady=10)

teach_notebook = ttk.Notebook(teach_tab)
teach_notebook.pack(expand=True, fill="both", padx=10, pady=10)

settings_tab = tk.Frame(teach_notebook, bg="#FFFBEA")
teach_notebook.add(settings_tab, text="⚙️ Wordlists & Settings")


phonics_map_tab = tk.Frame(teach_notebook, bg="#FFFBEA")
teach_notebook.add(phonics_map_tab, text="🔤 Phonics Sounds")
build_phonics_editor_gui(phonics_map_tab)

converter_tab = tk.Frame(teach_notebook, bg="#FFFBEA")
teach_notebook.add(converter_tab, text="📄 Build Phonics File")
select_phonics_btn = tk.Button(converter_tab, text="🎯 Set Phonics Wordlist", font=(CHOSEN_FONT, 12), bg="#F3E5F5", command=choose_phonics_list)
select_phonics_btn.pack(pady=5)



build_phonics_converter_gui(converter_tab)
build_word_fixer_gui(converter_tab)



manager_btn = tk.Button(settings_tab, text="⚙️ Word List Manager", font=(CHOSEN_FONT, 12), command=password_protected_manager)
manager_btn.pack(pady=5)

change_pw_btn = tk.Button(settings_tab, text="🔑 Change Admin Password",font=(CHOSEN_FONT, 10), command=change_admin_password)
change_pw_btn.pack(pady=5)

upload_btn = tk.Button(settings_tab, text="📂 Upload Words", font=(CHOSEN_FONT, 12), bg="#FFFACD", fg="#000", command=upload_word_file)
upload_btn.pack(pady=5)

active_list_label = tk.Label(settings_tab, text=f"📂 Current List: {config['active_list']}", font=(CHOSEN_FONT, 10), fg="#555", bg="#FFFBEA")
active_list_label.pack(pady=(5, 0))



#*******END of TEACH TAB**********


#---------------------------------


#*******LISTEN TAB**********

tk.Label(listen_tab, text="👋 Paste or type something below, then click to hear it!", font=(CHOSEN_FONT, 14), bg="#F0F8FF", fg="#333").pack(pady=(10, 5))

# Create a frame to hold both margin and text box
text_frame = tk.Frame(listen_tab, bg="#F0F8FF")
text_frame.pack(padx=20, pady=10, fill="both", expand=True)

# Left-side canvas for emoji line tracker
margin_canvas = tk.Canvas(text_frame, width=30, bg="#F0F8FF", highlightthickness=0)
margin_canvas.pack(side="left", fill="y")

# Main text box
text_box = tk.Text(text_frame, height=8, font=(CHOSEN_FONT, 18), wrap="word", bg="#FFFFFF", fg="#222", insertbackground="#000000")
text_box.pack(side="right", fill="both", expand=True)

#==Rate Slider==

toggle_btn = tk.Button(root, text="⚙️ Show Voice Settings", font=(CHOSEN_FONT, 11), command=toggle_settings)
toggle_btn.place(relx=0.97, rely=0.055, anchor="ne")  # upper right corner

settings_frame = tk.LabelFrame(root, text="Voice Settings", font=(CHOSEN_FONT, 12), bg="#F0F8FF", padx=10, pady=10)

disabled_label = tk.Label(settings_frame, text="(voice settings coming soon)", font=(CHOSEN_FONT, 9, "italic"), fg="#888", bg="#F0F8FF")
disabled_label.pack(pady=(0, 5))


rate_slider = tk.Scale(settings_frame, from_=100, to=250, label="Speed (wpm)", orient="horizontal", length=200)
rate_slider.set(150)
rate_slider.pack(pady=(5, 10))

preview_btn = tk.Button(settings_frame, text="🔊 Preview Voice", command=preview_voice)
preview_btn.pack()

read_style = tk.StringVar(value="standard")  # Options: standard or slow

style_label = tk.Label(settings_frame, text="Read Style:", font=(CHOSEN_FONT, 11), bg="#F0F8FF")
style_label.pack()

standard_radio = tk.Radiobutton(settings_frame, text="🎵 Standard", variable=read_style, value="standard", font=(CHOSEN_FONT, 10), bg="#F0F8FF")
standard_radio.pack(anchor="w")

slow_radio = tk.Radiobutton(settings_frame, text="🐢 Slow & Steady", variable=read_style, value="slow", font=(CHOSEN_FONT, 10), bg="#F0F8FF")
slow_radio.pack(anchor="w")

rate_slider.config(state="disabled")
preview_btn.config(state="disabled")
standard_radio.config(state="disabled")
slow_radio.config(state="disabled")
toggle_btn.config(state="disabled")

switch_user_btn = tk.Button(root, text="👤 Switch User", font=(CHOSEN_FONT, 11), command=open_profile_picker)
switch_user_btn.place(relx=0.85, rely=0.055, anchor="ne")  # top right



loading_label = tk.Label(root, text="🎤 Preparing voice...", font=(CHOSEN_FONT, 12), fg="#555", bg="#F0F8FF")


# === Button Row ===
button_frame = tk.Frame(listen_tab, bg="#F0F8FF")
button_frame.pack(pady=(0, 30))

speak_btn = tk.Button(button_frame, text="🎧 Read Aloud", font=(CHOSEN_FONT, 16, "bold"), bg="#4CAF50", fg="white", padx=20, pady=10, command=speak_text)
speak_btn.grid(row=0, column=0, padx=10)

replay_btn = tk.Button(button_frame, text="🔁 Replay", font=(CHOSEN_FONT, 16, "bold"), bg="#6495ED", fg="white", padx=20, pady=10, command=replay_last_text)
replay_btn.grid(row=0, column=2, padx=10)

pause_btn = tk.Button(button_frame, text="⏸ Pause", font=(CHOSEN_FONT, 16, "bold"), bg="#DAA520", fg="white", padx=20, pady=10, command=pause_audio)
pause_btn.grid(row=1, column=0, padx=10, pady=5)

resume_btn = tk.Button(button_frame, text="▶️ Resume", font=(CHOSEN_FONT, 16, "bold"), bg="#32CD32", fg="white", padx=20, pady=10, command=resume_audio)
resume_btn.grid(row=1, column=1, padx=10, pady=5)

stop_btn = tk.Button(button_frame, text="⏹ Stop", font=(CHOSEN_FONT, 16, "bold"), bg="#B22222", fg="white", padx=20, pady=10, command=stop_audio)
stop_btn.grid(row=1, column=2, padx=10, pady=5)

clear_btn = tk.Button(button_frame, text="🗑️ Clear Text", font=(CHOSEN_FONT, 16, "bold"), bg="#FF6347", fg="white", padx=20, pady=10, command=lambda: (text_box.delete("1.0", tk.END), text_box.tag_remove("sentence_highlight", "1.0", tk.END), margin_canvas.delete("ball")))
clear_btn.grid(row=0, column=1, padx=10)

#Right Click Read From Here

context_menu = tk.Menu(root, tearoff=0)
context_menu.add_command(label="📖 Read from here", command=read_from_cursor)
text_box.bind("<Button-3>", show_context_menu)  # Right-click for Windows/Linux
text_box.bind("<Button-2>", show_context_menu)  # Optional for macOS


#=============================================


#*******Write TAB**********



# Wrapper frame
container = tk.Frame(write_tab, bg="#FFFDF7")
container.pack(fill="both", expand=True, padx=20, pady=20)

# Left column for text box
left_frame = tk.Frame(container, bg="#FFFDF7")
left_frame.grid(row=0, column=0, sticky="n")

# Right column for controls
right_frame = tk.Frame(container, bg="#FFFDF7")
right_frame.grid(row=0, column=1, padx=20, sticky="n")

write_text = tk.Text(left_frame, font=(CHOSEN_FONT, 16), wrap="word", width=50, height=18, padx=10, pady=10, bg="#FFFFFF")
write_text.pack()

clear_btn = tk.Button(left_frame, text="🗑️ Clear Text", font=(CHOSEN_FONT, 12), bg="#FFC0CB", command=lambda: write_text.delete("1.0", tk.END))
clear_btn.pack(pady=(10, 0))

mic_btn = tk.Button(left_frame, text="🎤 Talk to Write", font=(CHOSEN_FONT, 12), bg="#DFF5FF", command=start_voice_input)
mic_btn.pack(pady=5)



test_voice_btn = tk.Button(left_frame, text="🔊 Test Voice", font=(CHOSEN_FONT, 11), command=test_typing_voice)
test_voice_btn.pack()




feedback_label = tk.Label(right_frame, text="", font=(CHOSEN_FONT, 12), bg="#FFFDF7")
feedback_label.pack(pady=5)

toggle_btn = tk.Button(right_frame, text="🧠 Toggle Dyslexic Font", font=(CHOSEN_FONT, 11), command=toggle_dyslexic)
toggle_btn.pack(pady=(10, 5))

tk.Button(right_frame, text="📂 Upload Wordlist", font=(CHOSEN_FONT, 11), command=upload_challenge_wordlist).pack(pady=5)



tk.Button(right_frame, text="🎧 Start Listen & Type", font=(CHOSEN_FONT, 12), command=start_listen_and_type_challenge).pack(pady=5)

tk.Button(right_frame, text="✅ Submit Answer", font=(CHOSEN_FONT, 12), command=check_challenge_answer).pack(pady=5)

tk.Button(right_frame, text="🔁 Replay Word", font=(CHOSEN_FONT, 12), command=lambda: speak_with_gtts(challenge_wordlist[challenge_index])).pack(pady=5)








def on_key_release(event):
    color_new_char(event)
    sound_out_word(event)

write_text.bind("<KeyRelease>", on_key_release)
write_text.bind("<Button-3>", speak_word)  # Right-click to speak word



#=============================================




#*******START OF TYPING GAME**********


current_typing_word = ""





def start_typing_game():
    clear_play_area()
    menu_frame.pack_forget()

    build_typing_game_tab()

    if not challenge_wordlist:
        messagebox.showwarning("⚠️ No words loaded", "Please upload a challenge wordlist first.")


   



def upload_new_typing_wordlist():
    global challenge_wordlist
    file_path = filedialog.askopenfilename(
        title="Select Wordlist",
        filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
    )
    if not file_path:
        return

    try:
        with open(file_path) as f:
            words = [line.strip().lower() for line in f if line.strip()]
        if not words:
            messagebox.showwarning("Empty File", "No words found in that file.")
            return

        challenge_wordlist = words
        messagebox.showinfo("✅ New Wordlist Loaded", f"{len(words)} words loaded!")
        status_label.config(text="📂 New list uploaded!", fg="#333")

        # Reload with the first word from new list
        load_next_typing_word()
    except Exception as e:
        messagebox.showerror("❌ Error", f"Could not load wordlist:\n{e}")



    

def build_typing_game_tab():
    global current_typing_word, typing_score, dyslexic_mode, challenge_wordlist

    current_typing_word = ""
    typing_score = {"attempts": 0, "correct": 0}
    praise_messages = [
        "Amazing!", "Great job!", "Fantastic!", "You're a star!", "Brilliant!", "Yes! That's perfect!"
    ]

    typing_tab = tk.Frame(play_tab, bg="#FDF6E3")
    typing_tab.pack(fill="both", expand=True)

    tk.Label(typing_tab, text="🎯 Type the Word!", font=(CHOSEN_FONT, 16), bg="#FDF6E3", pady=10).pack()

    prompt_label = tk.Label(typing_tab, text="🔊 Listen & Type:", font=(CHOSEN_FONT, 14),
                            fg="#333", bg="#FDF6E3")
    prompt_label.pack(pady=(5, 0))

    word_frame = tk.Frame(typing_tab, bg="#FDF6E3")
    word_frame.pack(pady=(0, 10))

    input_box = tk.Entry(typing_tab, font=(CHOSEN_FONT, 16), width=25)
    input_box.pack(pady=10)

    status_label = tk.Label(typing_tab, text="", font=(CHOSEN_FONT, 12), bg="#FDF6E3")
    status_label.pack(pady=5)

    def style_typing_word(word):
        for widget in word_frame.winfo_children():
            widget.destroy()

        font_choice = "OpenDyslexic" if dyslexic_mode else CHOSEN_FONT
        vowels = set("aeiou")

        for char in word.upper():
            color = "#1976D2" if char.lower() in vowels else "#D32F2F"
            lbl = tk.Label(word_frame, text=char, font=(font_choice, 20, "bold"),
                           fg=color, bg="#FDF6E3", padx=2)
            lbl.pack(side="left")

    def load_next_typing_word():
        global current_typing_word
        current_typing_word = random.choice(challenge_wordlist) if challenge_wordlist else "apple"
        style_typing_word(current_typing_word)
        input_box.delete(0, tk.END)
        speak_with_gtts(current_typing_word)


    def check_typing_word():
        typed = input_box.get().strip().lower()
        target = current_typing_word
        typing_score["attempts"] += 1
        input_box.delete(0, tk.END)

        if not typed:
            status_label.config(text="⚠️ Please type the word before submitting.", fg="#FF8F00")
            speak_with_gtts("Please type the word before submitting.")
            return

        if typed == target:
            typing_score["correct"] += 1
            praise = random.choice(praise_messages)
            status_label.config(text=f"✅ {praise}", fg="#2E7D32")
            speak_with_gtts(praise)
        else:
            status_label.config(text=f"❌ You typed: {typed}\n✅ Correct spelling: {target}", fg="#D32F2F")
            speak_with_gtts(f"You typed {typed}. The correct spelling is {target}")

        load_next_typing_word()


    def upload_new_typing_wordlist():
        file_path = filedialog.askopenfilename(
            title="Select Wordlist",
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        if not file_path:
            return

        try:
            with open(file_path) as f:
                words = [line.strip().lower() for line in f if line.strip()]
            if not words:
                messagebox.showwarning("Empty File", "No words found in that file.")
                return

            challenge_wordlist[:] = words
            messagebox.showinfo("✅ New Wordlist Loaded", f"{len(words)} words loaded!")
            status_label.config(text="📂 New list uploaded!", fg="#333")
            load_next_typing_word()
        except Exception as e:
            messagebox.showerror("❌ Error", f"Could not load wordlist:\n{e}")

    tk.Button(typing_tab, text="✅ Submit", font=(CHOSEN_FONT, 12), bg="#AED581", command=check_typing_word).pack(pady=5)
    tk.Button(typing_tab, text="🔁 Replay Word", font=(CHOSEN_FONT, 11), command=lambda: speak_with_gtts(current_typing_word)).pack(pady=5)
    tk.Button(typing_tab, text="📈 Show Score", font=(CHOSEN_FONT, 11), command=lambda: messagebox.showinfo(
        "Score", f"Correct: {typing_score['correct']}\nAttempts: {typing_score['attempts']}")).pack(pady=5)
    tk.Button(typing_tab, text="📂 Upload Wordlist", font=(CHOSEN_FONT, 11), bg="#FFF9C4", command=upload_new_typing_wordlist).pack(pady=(5, 10))
    tk.Button(typing_tab, text="🔙 Back to Menu", font=(CHOSEN_FONT, 11), bg="#FFDDDD", command=show_menu).pack(pady=(10, 5))

    load_next_typing_word()




#*******END OF TYPING GAME**********




#====================================================

#*******START OF ADDITION GAME**********

fixed_number_mode = False
fixed_number_value = None



DIFFICULTY_LEVELS = {
    1: {"min": 1, "max": 9},       # Level 1: Single-digit
    2: {"min": 10, "max": 99},     # Level 2: Two-digit
    3: {"min": 100, "max": 999}    # Level 3: Three-digit
}

def set_difficulty_level():
    if not check_password():
        messagebox.showwarning("Access Denied", "Incorrect password.")
        return

    level = simpledialog.askinteger("🎯 Select Level", "Choose difficulty level (1–3):")
    if level not in DIFFICULTY_LEVELS:
        messagebox.showerror("Invalid Level", "Please choose a level between 1 and 3.")
        return

    with open("difficulty_config.json", "w") as f:
        json.dump({"level": level}, f)

    messagebox.showinfo("✅ Level Set", f"Difficulty set to Level {level}.")


def load_selected_level():
    try:
        with open("difficulty_config.json", "r") as f:
            config = json.load(f)
            return config.get("level", 1)
    except FileNotFoundError:
        return 1  # Default to Level 1

def load_next_addition_problem():
    level = load_selected_level()
    bounds = DIFFICULTY_LEVELS.get(level, DIFFICULTY_LEVELS[1])
    a = random.randint(bounds["min"], bounds["max"])
    b = random.randint(bounds["min"], bounds["max"])
    return a, b



def start_addition_game():
    clear_play_area()
    menu_frame.pack_forget()

    build_addition_game_tab()


def build_addition_game_tab():
    global current_addition_problem, addition_score, timer_id
    clear_play_area()
    menu_frame.pack_forget()

    timer_id = None
    addition_score = {"attempts": 0, "correct": 0}
    praise_messages = [
        "Awesome!", "You're a math whiz!", "Great job!", "Correct!", "Brilliant!", "Nice work!"
    ]

    


    addition_tab = tk.Frame(play_tab, bg="#E3F2FD")
    addition_tab.pack(fill="both", expand=True)

    tk.Label(addition_tab, text="➕ Add It Up!", font=(CHOSEN_FONT, 16), bg="#E3F2FD", pady=10).pack()

    progress_bar = ttk.Progressbar(addition_tab, length=200, mode='determinate')
    progress_bar.pack(pady=5)
    progress_bar["maximum"] = 10  # Show progress toward next teacher check
    progress_bar["value"] = 0     # Start at 0

    # 🎯 Difficulty Selector
    level_var = tk.IntVar(value=load_selected_level())



    def update_level():
        selected = level_var.get()
        with open("difficulty_config.json", "w") as f:
            json.dump({"level": selected}, f)
        messagebox.showinfo("✅ Level Updated", f"Difficulty set to Level {selected}")
        load_problem_for_tab()

    
    tk.Label(addition_tab, text="🎯 Select Level:", font=(CHOSEN_FONT, 12), bg="#E3F2FD").pack()
    tk.OptionMenu(addition_tab, level_var, 1, 2, 3).pack()
    tk.Button(addition_tab, text="⚙️ Update Level", font=(CHOSEN_FONT, 11), command=update_level).pack(pady=5)
        

    # 🧮 Problem Prompt
    tk.Label(addition_tab, text="🧮 Solve this:", font=(CHOSEN_FONT, 14),
             fg="#333", bg="#E3F2FD").pack(pady=(5, 0))

    problem_frame = tk.Frame(addition_tab, bg="#E3F2FD")
    problem_frame.pack(pady=(0, 10))
    # Initialize this somewhere in your setup
    problem_frame.answered_count = 0


    input_box = tk.Entry(addition_tab, font=(CHOSEN_FONT, 16), width=10)
    input_box.bind("<Return>", lambda event: check_addition_answer())

    input_box.pack(pady=10)

    status_label = tk.Label(addition_tab, text="", font=(CHOSEN_FONT, 12), bg="#E3F2FD")
    status_label.pack(pady=5)

    timer_label = tk.Label(addition_tab, text="", font=(CHOSEN_FONT, 12), fg="#D84315", bg="#E3F2FD")
    timer_label.pack(pady=5)



    def activate_fixed_number_mode():
        global fixed_number_mode, fixed_number_value
        fixed = simpledialog.askinteger("🔢 Fixed Number Mode", "Enter a number to fix (e.g. 9):", minvalue=0, maxvalue=999)
        if fixed is None:
            return
        fixed_number_mode = True
        fixed_number_value = fixed
        messagebox.showinfo("🎯 Fixed Number Mode", f"Now practicing: {fixed} + x")
        load_problem_for_tab()

    #tk.Button(addition_tab, text="🔢 Fixed Number Mode", font=(CHOSEN_FONT, 11), bg="#FFF9C4", command=activate_fixed_number_mode).pack(pady=5)

    def deactivate_fixed_number_mode():
        global fixed_number_mode, fixed_number_value
        fixed_number_mode = False
        fixed_number_value = None
        messagebox.showinfo("🔚 Mode Off", "Fixed Number Mode has been turned off.")
        load_problem_for_tab()

    #tk.Button(addition_tab, text="❌ Exit Fixed Mode", font=(CHOSEN_FONT, 11), bg="#FFCDD2", command=deactivate_fixed_number_mode).pack(pady=5)

    
     # 📦 Create a horizontal frame for the buttons
    fixed_mode_frame = tk.Frame(addition_tab, bg="#E3F2FD")
    fixed_mode_frame.pack(pady=5)

    # 🔢 Fixed Number Mode Button
    tk.Button(fixed_mode_frame, text="🔢 Fixed Number Mode", font=(CHOSEN_FONT, 11),  bg="#FFF9C4", command=activate_fixed_number_mode).pack(side="left", padx=5)

    # ❌ Exit Fixed Mode Button
    tk.Button(fixed_mode_frame, text="❌ Exit Fixed Mode", font=(CHOSEN_FONT, 11),   bg="#FFCDD2", command=deactivate_fixed_number_mode).pack(side="left", padx=5)   
    
        

    def style_addition_problem(a, b):
        for widget in problem_frame.winfo_children():
            widget.destroy()

        answer = str(a + b)
        width = len(answer)

        a_str = str(a).rjust(width, '0')
        b_str = str(b).rjust(width, '0')

        # Create a grid frame
        grid_frame = tk.Frame(problem_frame, bg="#E3F2FD")
        grid_frame.pack(anchor="center", pady=10)

        carry_entries = []
        digit_entries = []

        # Row 0: Carry boxes (skip last digit)
        for i in range(width - 1):
            carry = tk.Entry(grid_frame, font=(CHOSEN_FONT, 14), width=2, justify="center")
            carry.grid(row=0, column=i + 1, padx=2, pady=2)
            carry_entries.append(carry)

        # Fill last column with empty space for alignment
        tk.Label(grid_frame, text="", bg="#E3F2FD", width=2).grid(row=0, column=width, padx=2)

        # Row 1: First number digits
        for i, digit in enumerate(a_str):
            tk.Label(grid_frame, text=digit, font=(CHOSEN_FONT, 20, "bold"),
                     fg="#1976D2", bg="#E3F2FD", width=2).grid(row=1, column=i + 1, padx=2)

        # Row 2: Second number digits with separate '+' sign
        tk.Label(grid_frame, text="+", font=(CHOSEN_FONT, 20, "bold"),
                 fg="#D32F2F", bg="#E3F2FD", width=2).grid(row=2, column=0, padx=2)

        for i, digit in enumerate(b_str):
            tk.Label(grid_frame, text=digit, font=(CHOSEN_FONT, 20, "bold"),
                     fg="#D32F2F", bg="#E3F2FD", width=2).grid(row=2, column=i + 1, padx=2)

        # Row 3: Divider line
        for i in range(width):
            tk.Label(grid_frame, text="―", font=(CHOSEN_FONT, 20),
                     fg="#333", bg="#E3F2FD", width=2).grid(row=3, column=i + 1)

        # Row 4: Answer input boxes
        for i in range(width):
            entry = tk.Entry(grid_frame, font=(CHOSEN_FONT, 20), width=2, justify="center")
            entry.grid(row=4, column=i + 1, padx=2)
            digit_entries.append(entry)

        # Store for later use
        problem_frame.digit_entries = digit_entries
        problem_frame.carry_entries = carry_entries


    def get_timer_duration(level):
        return {1: 20, 2: 30, 3: 45}.get(level, 10)

    def start_timer(seconds):
        global timer_id

        def countdown(secs_remaining):
            global timer_id
            if secs_remaining <= 0:
                timer_label.config(text="⏰ Time's up!")
                speak_with_gtts("Time has run out. Try again.")
                messagebox.showwarning("⏰ Time's Up", "Time has run out. Try again.")
                load_problem_for_tab()
                return
            timer_label.config(text=f"⏱️ Time left: {secs_remaining}s")
            timer_id = addition_tab.after(1000, lambda: countdown(secs_remaining - 1))

        countdown(seconds)

    def cancel_timer():
        global timer_id
        if timer_id:
            addition_tab.after_cancel(timer_id)
            timer_id = None



    def load_problem_for_tab():
        global current_addition_problem, fixed_number_mode

        cancel_timer()
        level = load_selected_level()
        bounds = DIFFICULTY_LEVELS[level]

        if fixed_number_mode:
            a = fixed_number_value
            b = random.randint(bounds["min"], bounds["max"])
        else:
            a = random.randint(bounds["min"], bounds["max"])
            b = random.randint(bounds["min"], bounds["max"])

        current_addition_problem = (a, b)
        style_addition_problem(a, b)
        input_box.delete(0, tk.END)
        speak_with_gtts(f"What is {a} plus {b}?")
        start_timer(get_timer_duration(level))


    def show_score():
        messagebox.showinfo("Score", f"Correct: {addition_score['correct']}\nAttempts: {addition_score['attempts']}")


    def show_teacher_check():
        result = messagebox.askyesno("Teacher Check", "10 questions have been answered.\nWould you like to review the score now?")
        if result:
            show_score()  # Or whatever function displays the score


    def check_addition_answer():
        typed = input_box.get().strip()
        input_box.delete(0, tk.END)
        cancel_timer()

        if not typed.isdigit():
            status_label.config(text="⚠️ Please enter a number.", fg="#FF8F00")
            speak_with_gtts("Please enter a number.")
            return

        a, b = current_addition_problem
        correct_answer = a + b
        user_answer = int(typed)
        addition_score["attempts"] += 1

        problem_frame.answered_count += 1
        progress_bar["value"] = problem_frame.answered_count % 10


        if problem_frame.answered_count % 10 == 0:
            show_teacher_check()

        if user_answer == correct_answer:
            addition_score["correct"] += 1
            praise = random.choice(praise_messages)
            status_label.config(text=f"✅ {praise}", fg="#2E7D32")
            speak_with_gtts(praise)
        else:
            status_label.config(text=f"❌ You answered: {user_answer}\n✅ Correct answer: {correct_answer}", fg="#D32F2F")
            speak_with_gtts(f"You answered {user_answer}. The correct answer is {correct_answer}")

        load_problem_for_tab()



    # 🎮 Button Row Frame
    button_row = tk.Frame(addition_tab, bg="#E3F2FD")
    button_row.pack(pady=5)

    # ✅ Submit Button
    tk.Button(button_row, text="✅ Submit", font=(CHOSEN_FONT, 12), bg="#AED581", command=check_addition_answer).pack(side="left", padx=5)

    # 🔁 Replay Problem Button
    tk.Button(button_row, text="🔁 Replay Problem", font=(CHOSEN_FONT, 11),command=lambda: speak_with_gtts(f"What is {current_addition_problem[0]} plus {current_addition_problem[1]}?")).pack(side="left", padx=5)

    # 📈 Show Score Button
    tk.Button(button_row, text="📈 Show Score", font=(CHOSEN_FONT, 11), command=lambda: messagebox.showinfo("Score", f"Correct: {addition_score['correct']}\nAttempts: {addition_score['attempts']}")).pack(side="left", padx=5)

    # 🔙 Back to Menu Button (keep this separate below)
    tk.Button(addition_tab, text="🔙 Back to Menu", font=(CHOSEN_FONT, 11), bg="#FFDDDD", command=lambda: [cancel_timer(), show_menu()]).pack(pady=(10, 5))


    load_problem_for_tab()




#*******END OF ADDITION GAME**********




#====================================================


#*******START OF SUBTRACTION GAME**********

fixed_number_mode = False
fixed_number_value = None


def load_next_subtraction_problem():
    level = load_selected_level()
    bounds = DIFFICULTY_LEVELS.get(level, DIFFICULTY_LEVELS[1])
    a = random.randint(bounds["min"], bounds["max"])
    b = random.randint(bounds["min"], bounds["max"])
    if b > a:
        a, b = b, a  # Ensure non-negative result
    return a, b

def start_subtraction_game():
    clear_play_area()
    menu_frame.pack_forget()
    build_subtraction_game_tab()

def build_subtraction_game_tab():
    global current_subtraction_problem, subtraction_score, timer_id
    clear_play_area()
    menu_frame.pack_forget()

    timer_id = None
    subtraction_score = {"attempts": 0, "correct": 0}
    praise_messages = [
        "Great subtraction!", "You're sharp!", "Correct!", "Nice work!", "Well done!", "Subtraction master!"
    ]

    subtraction_tab = tk.Frame(play_tab, bg="#E3F2FD")
    subtraction_tab.pack(fill="both", expand=True)

    tk.Label(subtraction_tab, text="➖ Subtraction Sprint", font=(CHOSEN_FONT, 16), bg="#E3F2FD", pady=10).pack()

    progress_bar = ttk.Progressbar(subtraction_tab, length=200, mode='determinate')
    progress_bar.pack(pady=5)
    progress_bar["maximum"] = 10
    progress_bar["value"] = 0

    level_var = tk.IntVar(value=load_selected_level())

    def activate_fixed_number_mode_subtraction():
        global fixed_number_mode, fixed_number_value
        fixed = simpledialog.askinteger("🔢 Fixed Number Mode", "Enter a number to fix (e.g. 20):", minvalue=0, maxvalue=999)
        if fixed is None:
            return
        fixed_number_mode = True
        fixed_number_value = fixed
        messagebox.showinfo("🎯 Fixed Number Mode", f"Now practicing: {fixed} − x")
        load_subtraction_problem()

    def deactivate_fixed_number_mode_subtraction():
        global fixed_number_mode, fixed_number_value
        fixed_number_mode = False
        fixed_number_value = None
        messagebox.showinfo("🔚 Mode Off", "Fixed Number Mode has been turned off.")
        load_subtraction_problem()


    
    def update_level():
        selected = level_var.get()
        with open("difficulty_config.json", "w") as f:
            json.dump({"level": selected}, f)
        messagebox.showinfo("✅ Level Updated", f"Difficulty set to Level {selected}")
        load_subtraction_problem()

    tk.Label(subtraction_tab, text="🎯 Select Level:", font=(CHOSEN_FONT, 12), bg="#E3F2FD").pack()
    tk.OptionMenu(subtraction_tab, level_var, 1, 2, 3).pack()
    tk.Button(subtraction_tab, text="⚙️ Update Level", font=(CHOSEN_FONT, 11), command=update_level).pack(pady=5)

    tk.Label(subtraction_tab, text="🧮 Solve this:", font=(CHOSEN_FONT, 14),
             fg="#333", bg="#E3F2FD").pack(pady=(5, 0))

    problem_frame = tk.Frame(subtraction_tab, bg="#E3F2FD")
    problem_frame.pack(pady=(0, 10))
    problem_frame.answered_count = 0

    input_box = tk.Entry(subtraction_tab, font=(CHOSEN_FONT, 16), width=10)
    input_box.bind("<Return>", lambda event: check_subtraction_answer())
    input_box.pack(pady=10)

    status_label = tk.Label(subtraction_tab, text="", font=(CHOSEN_FONT, 12), bg="#E3F2FD")
    status_label.pack(pady=5)

    timer_label = tk.Label(subtraction_tab, text="", font=(CHOSEN_FONT, 12), fg="#D84315", bg="#E3F2FD")
    timer_label.pack(pady=5)

    def style_subtraction_problem(a, b):
        for widget in problem_frame.winfo_children():
            widget.destroy()

        answer = str(a - b)
        width = len(answer)

        a_str = str(a).rjust(width, '0')
        b_str = str(b).rjust(width, '0')

        grid_frame = tk.Frame(problem_frame, bg="#E3F2FD")
        grid_frame.pack(anchor="center", pady=10)

        for i, digit in enumerate(a_str):
            tk.Label(grid_frame, text=digit, font=(CHOSEN_FONT, 20, "bold"),
                     fg="#C62828", bg="#E3F2FD", width=2).grid(row=1, column=i + 1, padx=2)

        tk.Label(grid_frame, text="-", font=(CHOSEN_FONT, 20, "bold"),
                 fg="#AD1457", bg="#E3F2FD", width=2).grid(row=2, column=0, padx=2)

        for i, digit in enumerate(b_str):
            tk.Label(grid_frame, text=digit, font=(CHOSEN_FONT, 20, "bold"),
                     fg="#AD1457", bg="#E3F2FD", width=2).grid(row=2, column=i + 1, padx=2)

        for i in range(width):
            tk.Label(grid_frame, text="―", font=(CHOSEN_FONT, 20),
                     fg="#333", bg="#E3F2FD", width=2).grid(row=3, column=i + 1)

        digit_entries = []
        for i in range(width):
            entry = tk.Entry(grid_frame, font=(CHOSEN_FONT, 20), width=2, justify="center")
            entry.grid(row=4, column=i + 1, padx=2)
            digit_entries.append(entry)

        problem_frame.digit_entries = digit_entries

    def get_timer_duration(level):
        return {1: 20, 2: 30, 3: 45}.get(level, 10)

    def start_timer(seconds):
        global timer_id

        def countdown(secs_remaining):
            global timer_id
            if secs_remaining <= 0:
                timer_label.config(text="⏰ Time's up!")
                speak_with_gtts("Time has run out. Try again.")
                messagebox.showwarning("⏰ Time's Up", "Time has run out. Try again.")
                load_subtraction_problem()
                return
            timer_label.config(text=f"⏱️ Time left: {secs_remaining}s")
            timer_id = subtraction_tab.after(1000, lambda: countdown(secs_remaining - 1))

        countdown(seconds)

    def cancel_timer():
        global timer_id
        if timer_id:
            subtraction_tab.after_cancel(timer_id)
            timer_id = None


    
    def load_subtraction_problem():
        global current_subtraction_problem
        cancel_timer()
        a, b = load_next_subtraction_problem()
        current_subtraction_problem = (a, b)
        style_subtraction_problem(a, b)
        input_box.delete(0, tk.END)
        speak_with_gtts(f"What is {a} minus {b}?")
        start_timer(get_timer_duration(load_selected_level()))
    


    '''
    def load_next_subtraction_problem():
        level = load_selected_level()
        bounds = DIFFICULTY_LEVELS.get(level, DIFFICULTY_LEVELS[1])

        if fixed_number_mode and fixed_number_value is not None:
            a = fixed_number_value
            b = random.randint(bounds["min"], min(a, bounds["max"]))
        else:
            a = random.randint(bounds["min"], bounds["max"])
            b = random.randint(bounds["min"], bounds["max"])
            if b > a:
                a, b = b, a  # Ensure non-negative result

        return a, b
    '''

    def load_next_subtraction_problem():
        level = load_selected_level()
        bounds = DIFFICULTY_LEVELS.get(level, DIFFICULTY_LEVELS[1])

        if fixed_number_mode and fixed_number_value is not None:
            a = fixed_number_value
            b = random.randint(bounds["min"], min(a, bounds["max"]))
        else:
            a = random.randint(bounds["min"], bounds["max"])
            b = random.randint(bounds["min"], bounds["max"])
            if b > a:
                a, b = b, a  # Ensure non-negative result

        return a, b


    # 📦 Fixed Number Mode Buttons Frame
    fixed_mode_frame = tk.Frame(subtraction_tab, bg="#E3F2FD")
    fixed_mode_frame.pack(pady=5)

    # 🔢 Activate Fixed Mode
    tk.Button(fixed_mode_frame, text="🔢 Fixed Number Mode", font=(CHOSEN_FONT, 11), bg="#FFF9C4",
              command=activate_fixed_number_mode_subtraction).pack(side="left", padx=5)

    # ❌ Deactivate Fixed Mode
    tk.Button(fixed_mode_frame, text="❌ Exit Fixed Mode", font=(CHOSEN_FONT, 11), bg="#FFCDD2",
              command=deactivate_fixed_number_mode_subtraction).pack(side="left", padx=5)



    def show_score():
        messagebox.showinfo("Score", f"Correct: {subtraction_score['correct']}\nAttempts: {subtraction_score['attempts']}")

    def show_teacher_check():
        result = messagebox.askyesno("Teacher Check", "10 questions have been answered.\nWould you like to review the score now?")
        if result:
            show_score()

    def check_subtraction_answer():
        typed = input_box.get().strip()
        input_box.delete(0, tk.END)
        cancel_timer()

        if not typed.isdigit():
            status_label.config(text="⚠️ Please enter a number.", fg="#FF8F00")
            speak_with_gtts("Please enter a number.")
            return

        a, b = current_subtraction_problem
        correct_answer = a - b
        user_answer = int(typed)
        subtraction_score["attempts"] += 1

        problem_frame.answered_count += 1
        progress_bar["value"] = problem_frame.answered_count % 10

        if problem_frame.answered_count % 10 == 0:
            show_teacher_check()

        if user_answer == correct_answer:
            subtraction_score["correct"] += 1
            praise = random.choice(praise_messages)
            status_label.config(text=f"✅ {praise}", fg="#2E7D32")
            speak_with_gtts(praise)
        else:
            status_label.config(text=f"❌ You answered: {user_answer}\n✅ Correct answer: {correct_answer}", fg="#D32F2F")
            speak_with_gtts(f"You answered {user_answer}. The correct answer is {correct_answer}")

        load_subtraction_problem()

  

    # 🎮 Button Row Frame
    button_row = tk.Frame(subtraction_tab, bg="#E3F2FD")
    button_row.pack(pady=5)

    # ✅ Submit Button
    tk.Button(button_row, text="✅ Submit", font=(CHOSEN_FONT, 12), bg="#FFCDD2",
              command=check_subtraction_answer).pack(side="left", padx=5)

    # 🔁 Replay Problem Button
    tk.Button(button_row, text="🔁 Replay Problem", font=(CHOSEN_FONT, 11),
              command=lambda: speak_with_gtts(f"What is {current_subtraction_problem[0]} minus {current_subtraction_problem[1]}?")).pack(side="left", padx=5)

    # 📈 Show Score Button
    tk.Button(button_row, text="📈 Show Score", font=(CHOSEN_FONT, 11),
              command=show_score).pack(side="left", padx=5)

    # 🔙 Back to Menu Button (keep this separate below)
    tk.Button(subtraction_tab, text="🔙 Back to Menu", font=(CHOSEN_FONT, 11), bg="#FFDDDD",
              command=lambda: [cancel_timer(), show_menu()]).pack(pady=(10, 5))


    load_subtraction_problem()


#*******END OF SUBTRACTION GAME**********

#====================================================

LEVEL_LABELS = {
    1: "1-digit × 1-digit",
    2: "2-digit × 1-digit",
    3: "3-digit × 1-digit",
    4: "2-digit × 2-digit",
    5: "3-digit × 2-digit"
}


#*******START OF MULTIPLICATION GAME**********
'''
def load_next_multiplication_problem():
    level = load_selected_level()
    bounds = DIFFICULTY_LEVELS.get(level, DIFFICULTY_LEVELS[1])
    a = random.randint(bounds["min"], bounds["max"])
    b = random.randint(bounds["min"], bounds["max"])
    return a, b
'''

def load_next_multiplication_problem():
    level = load_selected_level()

    if level == 1:
        a = random.randint(1, 9)
        b = random.randint(1, 9)
    elif level == 2:
        a = random.randint(10, 99)
        b = random.randint(1, 9)
    elif level == 3:
        a = random.randint(100, 999)
        b = random.randint(1, 9)
    elif level == 4:
        a = random.randint(10, 99)
        b = random.randint(10, 99)
    elif level == 5:
        a = random.randint(100, 999)
        b = random.randint(10, 99)
    else:
        # Default fallback
        a = random.randint(1, 9)
        b = random.randint(1, 9)

    return a, b


def start_multiplication_game():
    clear_play_area()
    menu_frame.pack_forget()
    build_multiplication_game_tab()

def build_multiplication_game_tab():
    global current_multiplication_problem, multiplication_score, timer_id
    clear_play_area()
    menu_frame.pack_forget()

    timer_id = None
    multiplication_score = {"attempts": 0, "correct": 0}
    praise_messages = [
        "Multiplication master!", "Great job!", "Correct!", "You're on fire!", "Brilliant!", "Nice work!"
    ]

    multiplication_tab = tk.Frame(play_tab, bg="#FFF3E0")
    multiplication_tab.pack(fill="both", expand=True)

    tk.Label(multiplication_tab, text="✖️ Multiply Mania", font=(CHOSEN_FONT, 16), bg="#FFF3E0", pady=10).pack()

    progress_bar = ttk.Progressbar(multiplication_tab, length=200, mode='determinate')
    progress_bar.pack(pady=5)
    progress_bar["maximum"] = 10
    progress_bar["value"] = 0

    level_var = tk.IntVar(value=load_selected_level())

    def update_level():
        selected = level_var.get()
        with open("difficulty_config.json", "w") as f:
            json.dump({"level": selected}, f)
        messagebox.showinfo("✅ Level Updated", f"Difficulty set to Level {selected}")
        load_multiplication_problem()

    tk.Label(multiplication_tab, text="🎯 Select Level:", font=(CHOSEN_FONT, 12), bg="#FFF3E0").pack()
    #tk.OptionMenu(multiplication_tab, level_var, 1, 2, 3).pack()
    tk.OptionMenu(multiplication_tab, level_var, *LEVEL_LABELS.keys()).pack()
    tk.Button(multiplication_tab, text="⚙️ Update Level", font=(CHOSEN_FONT, 11), command=update_level).pack(pady=2)

    tk.Label(multiplication_tab, text="🧮 Solve this:", font=(CHOSEN_FONT, 14),
             fg="#333", bg="#FFF3E0").pack(pady=(2, 0))

    problem_frame = tk.Frame(multiplication_tab, bg="#FFF3E0")
    problem_frame.pack(pady=(0, 5))
    problem_frame.answered_count = 0

    input_box = tk.Entry(multiplication_tab, font=(CHOSEN_FONT, 16), width=10)
    input_box.bind("<Return>", lambda event: check_multiplication_answer())

    input_box.pack(pady=10)

    status_label = tk.Label(multiplication_tab, text="", font=(CHOSEN_FONT, 12), bg="#FFF3E0")
    status_label.pack(pady=5)

    timer_label = tk.Label(multiplication_tab, text="", font=(CHOSEN_FONT, 12), fg="#D84315", bg="#FFF3E0")
    timer_label.pack(pady=5)



    def style_multiplication_problem(a, b):
        for widget in problem_frame.winfo_children():
            widget.destroy()

        a_str = str(a)
        b_str = str(b)
        num_rows = len(b_str)
        max_width = len(a_str) + len(b_str)  # Max possible digits in result

        grid_frame = tk.Frame(problem_frame, bg="#FFF3E0")
        grid_frame.pack(anchor="center", pady=10)

        carry_entries = []
        partial_entries = []
        final_entries = []

        # Row 0: Carry boxes (skip last digit)
        for i in range(max_width - 1):
            carry = tk.Entry(grid_frame, font=(CHOSEN_FONT, 14), width=2, justify="center")
            carry.grid(row=0, column=i + 1, padx=2, pady=2)
            carry_entries.append(carry)

        # Fill last column with empty space for alignment
        tk.Label(grid_frame, text="", bg="#FFF3E0", width=2).grid(row=0, column=max_width, padx=2)

        # Row 1: Top number (multiplicand)
        for i, digit in enumerate(a_str.rjust(max_width)):
            tk.Label(grid_frame, text=digit if digit != ' ' else '',
                     font=(CHOSEN_FONT, 20, "bold"), fg="#EF6C00", bg="#FFF3E0", width=2).grid(row=1, column=i)

        # Row 2: Bottom number (multiplier) with × sign
        tk.Label(grid_frame, text="×", font=(CHOSEN_FONT, 20, "bold"),
                 fg="#D84315", bg="#FFF3E0", width=2).grid(row=2, column=0)

        for i, digit in enumerate(b_str.rjust(max_width - 1)):
            tk.Label(grid_frame, text=digit if digit != ' ' else '',
                     font=(CHOSEN_FONT, 20, "bold"), fg="#D84315", bg="#FFF3E0", width=2).grid(row=2, column=i + 1)

        # Row 3: Divider line
        for i in range(max_width):
            tk.Label(grid_frame, text="―", font=(CHOSEN_FONT, 20),
                     fg="#333", bg="#FFF3E0", width=2).grid(row=3, column=i)

        # Rows 4 to n: Partial product entry boxes
        for row_index in range(num_rows):
            row_entries = []
            shift = row_index  # shift for place value
            for i in range(max_width):
                entry = tk.Entry(grid_frame, font=(CHOSEN_FONT, 20), width=2, justify="center")
                entry.grid(row=4 + row_index, column=i)
                if i < shift:
                    entry.insert(0, "")
                    entry.config(state="disabled", disabledbackground="#FFF3E0")
                row_entries.append(entry)
            partial_entries.append(row_entries)

        # Final answer row
        for i in range(max_width):
            entry = tk.Entry(grid_frame, font=(CHOSEN_FONT, 20), width=2, justify="center")
            entry.grid(row=4 + num_rows, column=i)
            final_entries.append(entry)

        # Store for later use
        problem_frame.carry_entries = carry_entries
        problem_frame.partial_entries = partial_entries
        problem_frame.final_entries = final_entries





    def get_timer_duration(level):
        return {1: 30, 2: 45, 3: 60, 4: 120, 5: 200}.get(level, 10)

    def start_timer(seconds):
        global timer_id

        def countdown(secs_remaining):
            global timer_id
            if secs_remaining <= 0:
                timer_label.config(text="⏰ Time's up!")
                speak_with_gtts("Time has run out. Try again.")
                messagebox.showwarning("⏰ Time's Up", "Time has run out. Try again.")
                load_multiplication_problem()
                return
            timer_label.config(text=f"⏱️ Time left: {secs_remaining}s")
            timer_id = multiplication_tab.after(1000, lambda: countdown(secs_remaining - 1))

        countdown(seconds)

    def cancel_timer():
        global timer_id
        if timer_id:
            multiplication_tab.after_cancel(timer_id)
            timer_id = None

    def load_multiplication_problem():
        global current_multiplication_problem
        cancel_timer()
        a, b = load_next_multiplication_problem()
        current_multiplication_problem = (a, b)
        style_multiplication_problem(a, b)
        input_box.delete(0, tk.END)
        speak_with_gtts(f"What is {a} times {b}?")
        start_timer(get_timer_duration(load_selected_level()))

    def show_score():
        messagebox.showinfo("Score", f"Correct: {multiplication_score['correct']}\nAttempts: {multiplication_score['attempts']}")

    def show_teacher_check():
        result = messagebox.askyesno("Teacher Check", "10 questions have been answered.\nWould you like to review the score now?")
        if result:
            show_score()

    def check_multiplication_answer():
        typed = input_box.get().strip()
        input_box.delete(0, tk.END)
        cancel_timer()

        if not typed.isdigit():
            status_label.config(text="⚠️ Please enter a number.", fg="#FF8F00")
            speak_with_gtts("Please enter a number.")
            return

        a, b = current_multiplication_problem
        correct_answer = a * b
        user_answer = int(typed)
        multiplication_score["attempts"] += 1

        problem_frame.answered_count += 1
        progress_bar["value"] = problem_frame.answered_count % 10

        if problem_frame.answered_count % 10 == 0:
            show_teacher_check()

        if user_answer == correct_answer:
            multiplication_score["correct"] += 1
            praise = random.choice(praise_messages)
            status_label.config(text=f"✅ {praise}", fg="#2E7D32")
            speak_with_gtts(praise)
        else:
            status_label.config(text=f"❌ You answered: {user_answer}\n✅ Correct answer: {correct_answer}", fg="#D32F2F")
            speak_with_gtts(f"You answered {user_answer}. The correct answer is {correct_answer}")

        load_multiplication_problem()


    
    # 🎮 Buttons in two columns
    button_frame = tk.Frame(multiplication_tab, bg="#FFF3E0")
    button_frame.pack(pady=10)

    tk.Button(button_frame, text="✅ Submit", font=(CHOSEN_FONT, 12), bg="#FFD54F",
              command=check_multiplication_answer).grid(row=0, column=0, padx=5, pady=5, sticky="ew")

    tk.Button(button_frame, text="🔁 Replay Problem", font=(CHOSEN_FONT, 11),
              command=lambda: speak_with_gtts(
                  f"What is {current_multiplication_problem[0]} times {current_multiplication_problem[1]}?"
              )).grid(row=0, column=1, padx=5, pady=5, sticky="ew")

    tk.Button(button_frame, text="📈 Show Score", font=(CHOSEN_FONT, 11),
              command=show_score).grid(row=1, column=0, padx=5, pady=5, sticky="ew")

    tk.Button(button_frame, text="🔙 Back to Menu", font=(CHOSEN_FONT, 11), bg="#FFDDDD",
              command=lambda: [cancel_timer(), show_menu()]).grid(row=1, column=1, padx=5, pady=5, sticky="ew")

    load_multiplication_problem()


    

#*******END OF MULTIPLICATION GAME**********

#====================================================


#*******START OF SUDOKU PUZZLE GAME**********

'''
def is_valid(board, row, col, num):
    for i in range(9):
        if board[row][i] == num or board[i][col] == num:
            return False
    start_row, start_col = 3 * (row // 3), 3 * (col // 3)
    for i in range(3):
        for j in range(3):
            if board[start_row + i][start_col + j] == num:
                return False
    return True

def solve_board(board):
    for row in range(9):
        for col in range(9):
            if board[row][col] == 0:
                nums = list(range(1, 10))
                random.shuffle(nums)
                for num in nums:
                    if is_valid(board, row, col, num):
                        board[row][col] = num
                        if solve_board(board):
                            return True
                        board[row][col] = 0
                return False
    return True

def generate_full_board():
    board = [[0 for _ in range(9)] for _ in range(9)]
    solve_board(board)
    return board

def remove_cells(board, difficulty="easy"):
    levels = {"easy": 35, "medium": 45, "hard": 55}
    cells_to_remove = levels.get(difficulty, 35)
    puzzle = [row[:] for row in board]
    while cells_to_remove > 0:
        row, col = random.randint(0,8), random.randint(0,8)
        if puzzle[row][col] != 0:
            puzzle[row][col] = 0
            cells_to_remove -= 1
    return puzzle

def start_sudoku_game():
    clear_play_area()
    menu_frame.pack_forget()



    sudoku_tab = tk.Frame(play_tab, bg="#F3E5F5")
    sudoku_tab.pack(fill="both", expand=True)

    tk.Label(sudoku_tab, text="🧩 Sudoku Challenge", font=(CHOSEN_FONT, 16), bg="#F3E5F5").pack(pady=10)

    difficulty = "medium"  # You can add a dropdown later
    full_board = generate_full_board()
    puzzle_board = remove_cells(full_board, difficulty)

    grid_frame = tk.Frame(sudoku_tab, bg="#F3E5F5")
    grid_frame.pack(pady=10)

    sudoku_entries = []

    for row in range(9):
        # Add horizontal divider after row 2 and 5
        if row in [3, 6]:
            divider_row = tk.Frame(grid_frame, height=2, bg="black")
            divider_row.grid(row=row * 2 - 1, column=0, columnspan=17, sticky="ew", pady=1)

        row_entries = []
        for col in range(9):
            val = puzzle_board[row][col]

            entry = tk.Entry(grid_frame, font=(CHOSEN_FONT, 14), width=2, justify="center", bd=1, relief="ridge")
            entry.grid(row=row * 2, column=col * 2, padx=1, pady=1)

            if val != 0:
                entry.insert(0, str(val))
                entry.config(state="disabled", disabledforeground="black")

            row_entries.append(entry)

            # Add vertical divider after column 2 and 5
            if col in [2, 5]:
                divider = tk.Frame(grid_frame, width=2, height=25, bg="black")
                divider.grid(row=row * 2, column=col * 2 + 1, sticky="ns", padx=1)

        sudoku_entries.append(row_entries)




    def check_solution():
        for r in range(9):
            for c in range(9):
                entry = sudoku_entries[r][c]
                val = entry.get()
                if val.isdigit() and int(val) == full_board[r][c]:
                    entry.config(bg="white")
                else:
                    entry.config(bg="#FFCDD2")  # Light red for incorrect

    tk.Button(sudoku_tab, text="✅ Check Solution", font=(CHOSEN_FONT, 12), bg="#CE93D8",
              command=check_solution).pack(pady=5)

    tk.Button(sudoku_tab, text="🔄 New Puzzle", font=(CHOSEN_FONT, 12), bg="#B3E5FC",
              command=start_sudoku_game).pack(pady=5)

    tk.Button(sudoku_tab, text="🔙 Back to Menu", font=(CHOSEN_FONT, 11), bg="#FFDDDD",
              command=show_menu).pack(pady=5)
'''



# Sudoku logic functions



# Sudoku logic functions
def is_valid(board, row, col, num):
    for i in range(9):
        if board[row][i] == num or board[i][col] == num:
            return False
    start_row, start_col = 3 * (row // 3), 3 * (col // 3)
    for i in range(3):
        for j in range(3):
            if board[start_row + i][start_col + j] == num:
                return False
    return True

def solve_board(board):
    for row in range(9):
        for col in range(9):
            if board[row][col] == 0:
                nums = list(range(1, 10))
                random.shuffle(nums)
                for num in nums:
                    if is_valid(board, row, col, num):
                        board[row][col] = num
                        if solve_board(board):
                            return True
                        board[row][col] = 0
                return False
    return True

def generate_full_board():
    board = [[0 for _ in range(9)] for _ in range(9)]
    solve_board(board)
    return board

def remove_cells(board, difficulty="easy"):
    levels = {"easy": 35, "medium": 45, "hard": 55}
    cells_to_remove = levels.get(difficulty, 35)
    puzzle = [row[:] for row in board]
    while cells_to_remove > 0:
        row, col = random.randint(0,8), random.randint(0,8)
        if puzzle[row][col] != 0:
            puzzle[row][col] = 0
            cells_to_remove -= 1
    return puzzle

# Global timer variables
start_time = None
timer_running = False
paused = False
pause_start = None
paused_duration = 0


def start_sudoku_game():
    global start_time, timer_running, paused, sudoku_entries, full_board, puzzle_board, difficulty_var

    clear_play_area()
    menu_frame.pack_forget()

    sudoku_tab = tk.Frame(play_tab, bg="#F3E5F5")
    sudoku_tab.pack(fill="both", expand=True)

    tk.Label(sudoku_tab, text="🧩 Sudoku Challenge", font=(CHOSEN_FONT, 16), bg="#F3E5F5").pack(pady=10)

    # Difficulty selector
    difficulty_frame = tk.Frame(sudoku_tab, bg="#F3E5F5")
    difficulty_frame.pack(pady=5)

    tk.Label(difficulty_frame, text="Difficulty:", font=(CHOSEN_FONT, 12), bg="#F3E5F5").pack(side="left", padx=5)

    difficulty_var = tk.StringVar(value="medium")
    difficulty_menu = ttk.Combobox(difficulty_frame, textvariable=difficulty_var, values=["easy", "medium", "hard"], state="readonly", width=10)
    difficulty_menu.pack(side="left", padx=5)

    # Timer label
    timer_label = tk.Label(sudoku_tab, text="Time: 00:00", font=(CHOSEN_FONT, 12), bg="#F3E5F5")
    timer_label.pack(pady=5)

    # Grid frame
    grid_frame = tk.Frame(sudoku_tab, bg="#F3E5F5")
    grid_frame.pack(pady=10)

    sudoku_entries = []

    def toggle_timer():
        global paused, pause_start, paused_duration, start_time
        paused = not paused
        if paused:
            pause_start = time.time()
            pause_button.config(text="▶ Resume")
        else:
            if pause_start:
                paused_duration += time.time() - pause_start
            pause_button.config(text="⏸ Pause")


    def update_timer():
        if timer_running and start_time:
            if not paused:
                elapsed = int(time.time() - start_time - paused_duration)
                mins, secs = divmod(elapsed, 60)
                timer_label.config(text=f"Time: {mins:02}:{secs:02}")
        timer_label.after(1000, update_timer)


    def load_new_puzzle():
        global full_board, puzzle_board, start_time, timer_running, paused
        difficulty = difficulty_var.get()
        full_board = generate_full_board()
        puzzle_board = remove_cells(full_board, difficulty)

        for widget in grid_frame.winfo_children():
            widget.destroy()
        sudoku_entries.clear()

        for row in range(9):
            if row in [3, 6]:
                divider_row = tk.Frame(grid_frame, height=2, bg="black")
                divider_row.grid(row=row * 2 - 1, column=0, columnspan=17, sticky="ew", pady=1)

            row_entries = []
            for col in range(9):
                val = puzzle_board[row][col]

                entry = tk.Entry(grid_frame, font=(CHOSEN_FONT, 14), width=2, justify="center", bd=1, relief="ridge")
                entry.grid(row=row * 2, column=col * 2, padx=1, pady=1)

                if val != 0:
                    entry.insert(0, str(val))
                    entry.config(state="disabled", disabledforeground="black")

                row_entries.append(entry)

                if col in [2, 5]:
                    divider = tk.Frame(grid_frame, width=2, height=25, bg="black")
                    divider.grid(row=row * 2, column=col * 2 + 1, sticky="ns", padx=1)

            sudoku_entries.append(row_entries)

        # Reset timer
        start_time = time.time()
        timer_running = True
        paused = False
        paused_duration = 0
        pause_start = None
        pause_button.config(text="⏸ Pause")

    def check_solution():
        incorrect_found = False
        for r in range(9):
            for c in range(9):
                entry = sudoku_entries[r][c]
                val = entry.get()
                if val.isdigit() and int(val) == full_board[r][c]:
                    entry.config(bg="white")
                else:
                    entry.config(bg="#FFCDD2")
                    incorrect_found = True

        if not incorrect_found:
            messagebox.showinfo("🎉 Success", "You successfully completed the Sudoku challenge!")

    # Buttons
    pause_button = tk.Button(sudoku_tab, text="⏸ Pause", font=(CHOSEN_FONT, 11), bg="#FFF9C4", command=toggle_timer)
    pause_button.pack(pady=5)

    tk.Button(sudoku_tab, text="✅ Check Solution", font=(CHOSEN_FONT, 12), bg="#CE93D8",
              command=check_solution).pack(pady=5)

    tk.Button(sudoku_tab, text="🔄 New Puzzle", font=(CHOSEN_FONT, 12), bg="#B3E5FC",
              command=load_new_puzzle).pack(pady=5)

    tk.Button(sudoku_tab, text="🔙 Back to Menu", font=(CHOSEN_FONT, 11), bg="#FFDDDD",
              command=show_menu).pack(pady=5)

    difficulty_menu.bind("<<ComboboxSelected>>", lambda e: load_new_puzzle())

    load_new_puzzle()
    update_timer()



#====================================================


#*******END OF SUDOKU PUZZLE GAME**********









#*******Game TAB BUTTONS**********




# === GAME TAB SETUP ===


# Shared title label
user = load_profiles().get("current_user", "Player")
game_label = tk.Label(play_tab, text=f"🎮 {user}'s Game Menu", font=(CHOSEN_FONT, 20), bg="#F9F9FF", fg="#333")
game_label.pack(pady=15)

# Game selection menu
menu_frame = tk.Frame(play_tab, bg="#F9F9FF")
menu_frame.pack()

profiles_data = load_profiles()
current_user = profiles_data.get("current_user", "Player")
name_label = tk.Label(menu_frame, text=f"👤 Logged in as: {current_user}", font=(CHOSEN_FONT, 10), fg="#444")
name_label.pack(pady=(0, 5))



# Game-specific widgets (defined later)
result_label = None
spell_frame = None

# Clear all game UI except menu
def clear_play_area():
    global spell_frame, result_label
    for widget in play_tab.winfo_children():
        if widget not in [game_label, menu_frame]:
            widget.destroy()
    spell_frame = None
    result_label = None
    user = load_profiles().get("current_user", "Player")
    game_label.config(text=f"🎮 {user}'s Game Menu")
    menu_frame.pack()


# Back to menu button


def show_menu():
    global after_ids  # or list each *_after_id individually if you're not using the dict

    # Cancel all scheduled callbacks
    for key in after_ids:
        if after_ids[key]:
            play_tab.after_cancel(after_ids[key])
            after_ids[key] = None

    clear_play_area()
    game_label.config(text=f"🎮 {load_profiles().get('current_user', 'Player')}'s Game Menu")
    menu_frame.pack()




def load_new_word():
    global spell_frame, result_label, back_btn, spell_words, session_results

    if result_label:
        result_label.config(text="")

    if not spell_words:
        messagebox.showerror("🚫 No Words Found", "The current word list is empty.\nPlease upload words or select a different list.")
        if back_btn and back_btn.winfo_exists():
            back_btn.config(state="normal")
        return

    if spell_frame:
        spell_frame.destroy()

    spell_frame = tk.Frame(play_tab, bg="#F9F9FF")
    spell_frame.pack()

    entry = random.choice(spell_words)
    correct = entry["word"]
    choices = random.sample(entry["choices"], len(entry["choices"]))

    # ✅ Load current user and profiles, if needed
    profiles = load_profiles()
    user = profiles.get("current_user")

    if back_btn and back_btn.winfo_exists():
        back_btn.config(state="disabled")

    def speak_and_enable():
        speak_with_gtts(f"Select the word {correct}")
        if back_btn and back_btn.winfo_exists():
            back_btn.config(state="normal")

    play_tab.after(300, speak_and_enable)

    # ✅ Now create the buttons
    for choice in choices:
        btn = tk.Button(spell_frame, text=choice, font=(CHOSEN_FONT, 14), command=lambda c=choice: check_spell_answer(c, correct), width=10)
        btn.pack(pady=5)

# Check answer logic

def check_spell_answer(selected, correct):
    global result_label

    profiles_data = load_profiles()
    user = profiles_data.get("current_user")

    if not user:
        print("⚠️ No user profile loaded.")
        return

    user_profile = profiles_data["profiles"].get(user, {})
    progress = user_profile.setdefault("progress", {})
    progress["games_played"] = progress.get("games_played", 0) + 1

    if selected == correct:
        progress["correct"] = progress.get("correct", 0) + 1
        if result_label:
            result_label.config(text="🎉 You got it!", fg="#4CAF50")
        speak_with_gtts("Awesome! That's correct.")
    else:
        progress["incorrect"] = progress.get("incorrect", 0) + 1
        if result_label:
            result_label.config(text="😬 Oops, try again!", fg="#E53935")
        speak_with_gtts("That's not it. Try another one.")

    # Save the updated progress
    save_profiles(profiles_data)
    check_milestones(user)


    # Load a new word after delay
    after_ids["spell_game"] = play_tab.after(2000, load_new_word)




def start_spell_game():
    global result_label, back_btn

    clear_play_area()
    menu_frame.pack_forget()
    game_label.config(text="🧠 Select the Word!")

    # Result Label
    result_label = tk.Label(play_tab, text="", font=(CHOSEN_FONT, 14), bg="#F9F9FF")
    result_label.pack(pady=10)

    # Back Button
    back_btn = tk.Button(
        play_tab,
        text="🔙 Back to Menu",
        font=(CHOSEN_FONT, 12),
        bg="#FFDDDD",
        command=show_menu,
        state="disabled"
    )
    back_btn.pack(pady=(0, 10))

    # ✅ Filter words that are valid for this game
    valid_words = [w for w in spell_words if "choices" in w and isinstance(w["choices"], list)]

    if not valid_words:
        result_label.config(text="❌ No compatible spelling entries found.\nUse Wordlist Manager or Upload Words.")
        back_btn.config(state="normal")

        return

    # ✅ Pick a clean list for use in the game
    global spell_game_words
    spell_game_words = valid_words
    load_new_word()









'''
# 🎲 Section Label
tk.Label(menu_frame, text="🎲 Fun Games", font=(CHOSEN_FONT, 14, "bold")).pack(pady=(20, 5))

# 🎯 Word Match (placeholder)
word_match_btn = tk.Button(menu_frame, text="🎯 Word Match", font=(CHOSEN_FONT, 13), bg="#FFF1D0", command=start_word_match)
word_match_btn.pack(pady=5)

# ⏱️ Quick Quiz (placeholder)
quick_quiz_btn = tk.Button(menu_frame, text="⏱️ Quick Quiz", font=(CHOSEN_FONT, 13), bg="#FFE4E1", command=start_quick_quiz)
quick_quiz_btn.pack(pady=5)
'''






# Create a container frame for both labs
labs_frame = tk.Frame(menu_frame)
labs_frame.pack(pady=(15, 5))

# 🎓 Language Lab Section
language_lab_frame = tk.Frame(labs_frame, width=200, height=300)
language_lab_frame.grid(row=0, column=0, padx=20, sticky="n")
language_lab_frame.grid_propagate(False)  # Prevent resizing

tk.Label(language_lab_frame, text="🎓 Language Lab", font=(CHOSEN_FONT, 20, "bold")).pack(pady=(0, 10))

tk.Button(language_lab_frame, text="▶️ Select the Word", font=(CHOSEN_FONT, 13), bg="#DFFFE2", command=start_spell_game).pack(pady=5)
tk.Button(language_lab_frame, text="🔊 Audio Challenge", font=(CHOSEN_FONT, 13), bg="#E6F4FF", command=start_audio_game).pack(pady=5)
tk.Button(language_lab_frame, text="🧩 Syllable Snap!", font=(CHOSEN_FONT, 13), bg="#E0F7FA", command=start_syllable_snap).pack(pady=5)
tk.Button(language_lab_frame, text="🔤 Sound It Out! (Coming Soon)", font=(CHOSEN_FONT, 13), bg="#FFF2E9", command=start_phonics_game, state="disabled").pack(pady=5)
tk.Button(language_lab_frame, text="🎯 Type the Word!", font=(CHOSEN_FONT, 13), bg="#FFF3E0", command=start_typing_game).pack(pady=5)

# 🎓 Math Lab Section
math_lab_frame = tk.Frame(labs_frame, width=200, height=300)
math_lab_frame.grid(row=0, column=1, padx=20, sticky="n")
math_lab_frame.grid_propagate(False)

tk.Label(math_lab_frame, text="🎓 Math Lab", font=(CHOSEN_FONT, 20, "bold")).pack(pady=(0, 10))
tk.Button(math_lab_frame, text="➕ Add It Up!", font=(CHOSEN_FONT, 13), bg="#E8F5E9", command=build_addition_game_tab).pack(pady=5)
tk.Button(math_lab_frame, text="➖ Subtraction Sprint", font=(CHOSEN_FONT, 13), bg="#FFEBEE", command=start_subtraction_game).pack(pady=5)
tk.Button(math_lab_frame, text="✖️ Multiply Mania", font=(CHOSEN_FONT, 13), bg="#FFEBEE", command=start_multiplication_game).pack(pady=5)
tk.Button(math_lab_frame, text="✖️ 🧩 Sudoku Challenge", font=(CHOSEN_FONT, 13), bg="#FFEBEE", command=start_sudoku_game).pack(pady=5)



'''
tk.Button(math_lab_frame, text="🧮 Number Ninja", font=(CHOSEN_FONT, 13), bg="#FFFDE7", command=start_number_ninja).pack(pady=5)
'''






def launch_profile_picker():
    profiles_data = load_profiles()

    if not profiles_data["profiles"]:
        # No profiles? Force creation
        messagebox.showinfo("👋 Welcome", "Let's create your first profile!")
        open_profile_picker()
    else:
        open_profile_picker()

root.after(100, launch_profile_picker)




# Play welcome message 1 second after GUI appears

threading.Thread(target=tts_worker, daemon=True).start()
root.mainloop()
